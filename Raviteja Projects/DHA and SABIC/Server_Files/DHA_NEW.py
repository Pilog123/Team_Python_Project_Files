import docx2pdf
from docx2pdf import convert
import sys
from subprocess import  Popen
# import cx_Oracle
from fastapi import FastAPI
from fastapi.params import Form
from fastapi.responses import FileResponse
from starlette.routing import Host
from starlette.responses import Response
import uvicorn
from uvicorn.main import run
from fastapi.responses import StreamingResponse
import subprocess
import pandas as pd
import docx
from docx import Document
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.shared import RGBColor,Pt,Inches
from docx.oxml import OxmlElement, ns
import warnings
warnings.filterwarnings("ignore")
import re
import threading
from datetime import datetime,date

from PyPDF2 import PdfFileMerger, PdfFileReader, PdfFileWriter
import cx_Oracle

import os
import win32com.client

wdFormatPDF = 17

app = FastAPI()


def convert_to_pdf_win(sp, dp):
    convert(sp, dp)
    return 'Done PDF Creation!'


def DHA_Caller(dha_data, RD):
    conn = cx_Oracle.connect("DR1024213/Pipl#mdrm$213@172.16.1.61/DR101413")
    BATCH_ID = RD['BATCH_ID']

    dha_data.rename(columns={'LONG_TEXT': 'LONG_DESCRIPTION', 'SHORT_TEXT': 'SHORT_DESCRIPTION', 'RECORD_NO': 'MATERIAL'}, inplace=True)
    dha_data = dha_data[['MATERIAL', 'LONG_DESCRIPTION', 'SHORT_DESCRIPTION', 'UOM']]
    dha_data.fillna('', inplace=True)
    today = date.today()
    date_ = today.strftime("%d")+"-"+today.strftime("%m")+"-"+today.strftime("%Y")
    # datetime object containing current date and time
    now = datetime.now()
    print('start time',now)
    
    insert_time = str(datetime.now().replace(microsecond=0))

    doc = docx.Document()
    # conn = cx_Oracle.connect("DR1024193/Pipl#mdrm$93@172.16.1.61/DR101412")

    # first page of document
    doc.add_paragraph('\n\n')
    #my_image = doc.add_picture('home/python_user/DOC/pi.jpg', width=Inches(2.79),height=Inches(0.90))
    my_image = doc.add_picture(r'C:\Users\Administrator\Desktop\pilog\pi.jpg', width=Inches(2.79),height=Inches(0.90))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
    doc.add_paragraph("\n\n")
    para = doc.add_paragraph().add_run('\tData Health Assessment')
    par = doc.add_paragraph().add_run('\t\t\t\t\t  For\n')
    parg = doc.add_paragraph().add_run('\t\t\tSAUDI CEMENT\n')
    pag = doc.add_paragraph().add_run('\t\t    Document Number: PDQH-ORG-DHA-001')
    para.bold = True
    parg.bold = True
    para.font.size = Pt(30)
    par.font.size = Pt(30)
    parg.font.size = Pt(30)
    pag.font.size = Pt(15)
    doc.add_page_break()

    # second page
    table = doc.add_table(rows=4,cols=4, style='Table Grid')
    A = table.cell(0,0)
    B = table.cell(0,1)
    D = table.cell(0,2)
    F = table.cell(0,3)
    C = A.merge(B)
    E = A.merge(D)
    G = A.merge(F)
    BB = table.cell(3,1)
    DD = table.cell(3,2)
    EE = table.cell(3,3)
    C = BB.merge(DD)
    F = BB.merge(EE)
    row = table.rows[0].cells
    row1 = table.rows[1].cells
    row2 = table.rows[2].cells
    row3 = table.rows[3].cells
    a=row[0].add_paragraph('Document Information')
    ab=row1[0].add_paragraph('Number')
    zz=row1[1].add_paragraph('PDQH-ORG-DHA-001')
    af=row1[2].add_paragraph('Version')
    Z1=row1[3].add_paragraph('1')
    ac=row2[0].add_paragraph('Effective')
    aT=row2[1].add_paragraph('Date: {}'.format(date_))
    ae=row2[2].add_paragraph('Issued')
    ap=row2[3].add_paragraph('Date: {}'.format(date_))
    ad=row3[0].add_paragraph('Compiled by')
    al=row3[1].add_paragraph('PiLog Cloud')
    a.bold = True
    ab.bold = True
    ac.bold = True
    ad.bold = True
    a.alignment=WD_TABLE_ALIGNMENT.CENTER
    intt = doc.add_paragraph().add_run('                        ')
    table = doc.add_table(rows=7,cols=5, style='Table Grid')
    row = table.rows[0].cells
    row11 = table.rows[1].cells
    row12 = table.rows[2].cells
    a=row[0].add_paragraph('Amendment History')
    a.alignment=WD_TABLE_ALIGNMENT.CENTER
    ab=row11[0].add_paragraph('Version')
    ab=row11[1].add_paragraph('Date')
    ab=row11[2].add_paragraph('Changed Chapter/Topic/Page')
    ab=row11[3].add_paragraph('Pages')
    ab=row11[4].add_paragraph('Checked By')
    zb=row12[0].add_paragraph('01')
    zc=row12[1].add_paragraph('Date: {}'.format(date_))
    zd=row12[2].add_paragraph('First Release')
    zde=row12[3].add_paragraph('17')
    zdu=row12[4].add_paragraph('CoE Team')
    A = table.cell(0,0)
    B = table.cell(0,1)
    D = table.cell(0,2)
    F = table.cell(0,3)
    H = table.cell(0,4)
    C = A.merge(B)
    E = A.merge(D)
    G = A.merge(F)
    I = A.merge(H)
    intt = doc.add_paragraph().add_run('                        ')
    doc.add_page_break()

    # About Document page
    line13 = doc.add_paragraph().add_run("""About This Document""")
    line13.font.size = Pt(18)
    line13.bold = True
    intt = doc.add_paragraph().add_run('''This document contains detailed information about the material master data of SAUDI CEMENT. An analysis was done on the data obtained from SAUDI CEMENT and the details of the analysis were compiled in this report.
    \nThis report contains the following information:\n
    The Scope of Analysis
    Methodology adapted
    Current status of SAUDI CEMENT data
    The next course of Action
    Suggestions and Recommendations''')
    doc.add_page_break()

    # Acronyms and Definitions
    Acronyms = doc.add_paragraph().add_run('Acronyms and Definitions')
    Acronyms.font.size = Pt(16)
    Acronyms.bold = True
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    row = table.rows[0].cells
    Acronym = ['ISO','ERP','FFT','IT','MDM','MDRM','MFR','OEM','OTD','PPO','PPR','PTNO','SPL','DQH']
    Definition = ['International Organization for Standardization','Enterprise Resource Planning','Free Format Text','Information Technology','Master Data Management','Master Data Record Manager','Manufacturer','Original Equipment Manufacturer','Open Technical Dictionary','PiLog Preferred Ontology','PiLog Preferred Records','Part Number','Supplier','Data Quality Hub']
    L5 = list(zip(Acronym,Definition))
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    pp=row[0].add_paragraph('Acronym')
    pp.bold = True
    pp.alignment=WD_TABLE_ALIGNMENT.CENTER
    p=row[1].add_paragraph('Definition')
    p.bold = True
    p.alignment=WD_TABLE_ALIGNMENT.CENTER
    for Acronym, Definition in zip(Acronym,Definition):
        row = table.add_row().cells
        table.cell(Acronym.index(Acronym),0).width =997280
        # row[0].text = str(Acronym)
        p1=row[0].add_paragraph(str(Acronym))
        p1.alignment=WD_TABLE_ALIGNMENT.CENTER
        table.cell(Definition.index(Definition),1).width =3846320
        p=row[1].add_paragraph(str(Definition))
        p.alignment=WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()

    # contents page
    con = doc.add_paragraph().add_run('\t                                                      Contents')
    con.font.color.rgb = RGBColor(0, 55, 140)
    con.bold = True
    con.font.size = Pt(14)
    coo = doc.add_paragraph().add_run("""Acronyms and Definitions..............................................................................................……………….4""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""1. PREFACE.................................................................................................………………...........................7""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""2. PROCESS & METHODOLOGY TO DELIVER SCOPE…………………………..…………….…………..7""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""3. CONSIDERATIONS ON DATA ANALYSIS………………………………………..………………….…….15""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""   3.1. Consideration on Data Analysis for DHA………………………………..………………………..15""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""4. DATA HEALTH AND QUALITY ANALYSIS RESULTS.................................………………...……17""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""   4.1. Detailed Data Analysis on Data Uniqueness……………………….………………………..…..17""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.1.1 Analysis on Material Records.......................…...……..…………………………...…………..17""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.1.2 Analysis on Descriptions/Texts.............….....................................................….…..……18""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.1.3 Advanced Analysis on Descriptions/Texts.......................................…..........….…..…20""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.1.4 Analysis on Reference Details...................…..………………………….................…….……22""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.1.5 Potential Duplicates w.r.t Manufacturer Part Number......….......….....…..….…..…23""")
    coo.bold = True

    coo = doc.add_paragraph().add_run("""   4.2 Detailed Data Analysis on Data Completeness.....…...................…………........….…...…..32""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.2.1 Analysis on Description Length..................................………...............................…..…..32""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.2.2 Report on Reference Details..............………..........................…………………………….....35""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""   4.3 Detailed Data Analysis on Data Consistency.......……….............................………...………36""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.3.1 Detailed Data Analysis on Data Consistency.......………..............….....………...………38""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.3.2 Non-standardized Prefixes in Description.........…..…............................……………...38""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""      4.3.3 Inconsistencies in UOMs.........................….............................……………………………..…38""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""   4.4 Analysis on Top 10 UOMs................................………………………............…........……………..39""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""   4.5 Analysis on Top 10 COMMODITYs................................………………………..........…..…..…..41""")
    coo.bold = True
    coo = doc.add_paragraph().add_run("""5. Key Recommendations...........................…………………..................……………...…............………42""")
    coo.bold = True
    # doc.add_page_break()

    # Preface page start
    cat1 = doc.add_paragraph().add_run("""1. PREFACE:""")
    cat1.font.size = Pt(16)
    cat1.bold = True
    intt = doc.add_paragraph().add_run("""Data is of high quality, if it fits for its intended uses in operations, decision-making, and planning.""")
    intt = doc.add_paragraph().add_run("""Data quality refers to the state of qualitative and quantitative pieces of information. Data is that which provides information about other data.""")
    intt = doc.add_paragraph().add_run("""A Data Health Assessment (DHA) profiles and analyzes the quality and integrity of your master data (e.g., Material, Vendor, Service, Asset, Customer, Business Partner, Equipment's, Location etc.,). A DHA visualizes,'AS-IS' state of your data quality, completeness, consistency, conformity to standards, and duplicates. The report then summarizes the impact that poor data has on your business and recommends the next steps to achieve measurable data and business metric optimization. The analysis process can be conducted across multiple domains that may involve datasets for customers, vendors, materials, and finance types of information""")
    cat2 = doc.add_paragraph().add_run("""2. PROCESS & METHODOLOGY TO DELIVER SCOP""")
    cat2.font.size = Pt(16)
    cat2.bold = True
    intt = doc.add_paragraph().add_run("""PiLog utilized proprietary automated processes to generate highly probable data points that indicate the data health and quality of the customer Data Set.""")
    intt = doc.add_paragraph().add_run("""These data quality processes and methodologies have been developed and proven over several years and allow us to readily identify certain characteristics indicating problem areas in both data and systems within an organization""")
    intt = doc.add_paragraph().add_run("""The approach is to determine records that are duplicated and unique, structured, standardized and rich in item property or characteristic values to give an overview of quality""")
    intt = doc.add_paragraph().add_run("""To determine the scope of the underlying root causes and to plan the ways that tools can be used to address data quality issues, it is valuable to understand these common data quality dimensions:""")
    doc.add_picture(r'C:\Users\Administrator\Desktop\SEP\METHODOLOGY.png', width=Inches(6), height=Inches(4))
    cat3 = doc.add_paragraph().add_run("""A. Completeness of Master Data:""")
    cat3.font.size = Pt(12)
    cat3.bold = True
    intt = doc.add_paragraph().add_run("""Data completeness refers to an indication of whether or not all the data necessary to meet the current and future business information demand are available in the data resource.""")
    intt = doc.add_paragraph().add_run("""It deals with determining the data needed to meet the business information demand and ensuring those data are captured and maintained in the data resource so they are available when needed. Various processes include Data Extraction, Data Transformation, Data Loading, Security implementation & Job Control.""")
    cat4 = doc.add_paragraph().add_run("""B. Redundancy of Master Data:""")
    cat4.font.size = Pt(12)
    cat4.bold = True
    intt = doc.add_paragraph().add_run("""Data redundancy ensures that data is not duplicated un-necessarily across any part of the system""")
    intt = doc.add_paragraph().add_run("""The online store may have a sales department and a complaints department. When a user buys a new product, this will be store as a transaction, but both the sales department and the complaints department may need this information at some point.""")
    intt = doc.add_paragraph().add_run("""It would be possible for each department to have its own separate database and each time there is a transaction it gets put into each database. This solution is not ideal though, as having to input the same data into multiple places risks user entry errors leading to data consistency problems.""")
    cat5 = doc.add_paragraph().add_run("""C. Compliance of Master Data:""")
    cat5.font.size = Pt(12)
    cat5.bold = True
    intt = doc.add_paragraph().add_run("""Data compliance refers to a state of being in accordance with established guidelines or specifications.""")
    intt = doc.add_paragraph().add_run("""The growing number of quality standards and regulations (industry specific or not) has also drawn attention to Master data management. In order to comply with these requirements, companies must meet certain criteria which are directly or indirectly impacted by the quality of data in the systems. There are many compliance risks that companies run from having bad Master data management:""")
    intt = doc.add_paragraph(style='List Bullet 2').add_run("""SOX risks occur in maintaining reporting structures and processing critical master data such as vendor bank accounts, fixed-asset data, contracts and contract conditions""")
    intt = doc.add_paragraph(style='List Bullet 2').add_run("""Petrochemical industry companies that are regulated by refining safety and operational standards and recommended practices may have significant exposure to legal risk and could even lose their operating licenses if their master records are incorrect with respect to product composition, storage locations, recording of ingredients, etc.""")
    intt = doc.add_paragraph(style='List Bullet 2').add_run("""Fiscal liabilities, such as VAT, produce risk. The VAT remittance may be incorrect if the relevant fields in the master data are not appropriately managed, possibly leading to inaccurate VAT percentages on intercompany sales""")
    cat6 = doc.add_paragraph().add_run("""D. Consistency & Integrity of Master Data:""")
    cat6.font.size = Pt(12)
    cat6.bold = True
    intt = doc.add_paragraph().add_run("""Data consistency refers to the transparency of the information, or the ability for other to see the changes and trends of data""")
    intt = doc.add_paragraph().add_run("""It is when you ensure that the same data that is being used in different parts of the system will always be the same. The data is consistent across the entire system""")
    intt = doc.add_paragraph().add_run("""In many examples, you may find the same data in needed in more than one place. In the online store example, the users address may be part of their signup information, but it will also need to appear on delivery information. Having to input the data directly into every place that it is needed causes problems if that data needs to be changed""")
    intt = doc.add_paragraph().add_run("""If the user moves home, then every place their address appears will need to be manually changed to ensure the data in consistent. If you do not do all of these updates, then you’ll have a situation where the address in one part of the system will be different to the address in another part of the system. This is solved by centralizing the data so there is one place that the address is stored. Any part of the system that needs to have the address can then just reference back to that central address location and find to find the data. When you do this you only have the change the data once for that change to spread through the whole system, ensuring that data always remains consistent.""")
    intt = doc.add_paragraph().add_run("""Data integrity refers to the accuracy and reliability of the data being collected""")
    cat7 = doc.add_paragraph().add_run("""E. Accuracy of Master Data:""")
    cat7.font.size = Pt(12)
    cat7.bold = True
    intt = doc.add_paragraph().add_run("""Data accuracy is one of the components of data quality. It refers to whether the data values stored for an object are the correct values. To be correct, a data values must be the right value and must be represented in a consistent and unambiguous form""")
    intt = doc.add_paragraph().add_run("""For example, birth date is December 13, 1941. If a personnel database has a BIRTH_DATE data element that expects dates in USA format, a date of 12/13/1941 would be correct. A date of 12/14/1941 would be inaccurate because it is the wrong value. A date of 13/12/1941 would be wrong because it is a European representation instead of a USA representation""")
    intt = doc.add_paragraph().add_run("""There are two characteristics of accuracy: form and content. Form is important because it eliminates ambiguities about the content. The birth date example is ambiguous because the reviewer would not know whether the date was invalid or just erroneously represented. In the case of a date such as 5 February, 1944, the USA representation is 02/05/1944, whereas the European representation is 05/02/1944. You cannot tell the representation from the value and thus need discipline in creating the date values in order to be accurate. A value is not accurate if the user of the value cannot tell what it is""")
    intt = doc.add_paragraph().add_run("""The concept of accuracy also applies above the data element level. Data elements are never recorded in isolation. They are value attributes of business objects such as personnel records, orders, invoices, payments, and inventory records. The business objects represent real-world objects or events, and each consists of one or more rows of one or more tables connected through keys. Object-level inaccuracies consist of objects that are missing, have missing parts, or that exist but should not""")
    cat8 = doc.add_paragraph().add_run("""F. Provenance of Master Data:""")
    cat8.font.size = Pt(12)
    cat8.bold = True
    intt = doc.add_paragraph().add_run("""Data provenance documents the inputs, entities, systems, and processes that influence data of interest, in effect providing a historical record of the data and its origins. The generated evidence supports essential forensic activities such as data-dependency analysis, error/compromise detection and recovery, and auditing and compliance analysis""")
    intt = doc.add_paragraph().add_run("""The provenance of data which is generated by complex transformations such as workflows is of considerable value. From it, one can ascertain the quality of the data based on its ancestral data and derivations, track back sources of errors, allow automated re-enactment of derivations to update a data, and provide attribution of data sources. Provenance is also essential to the business domain where it can be used to drill down to the source of data in a data warehouse, track the creation of intellectual property, and provide an audit trail for regulatory purposes""")
    intt = doc.add_paragraph().add_run("""The use of data provenance is proposed in distributed systems to trace records through a dataflow, replay the dataflow on a subset of its original inputs and debug data flows. To do so, one needs to keep track of the set of inputs to each operator, which were used to derive each of its outputs""")
    cat9 = doc.add_paragraph().add_run("""G. Uniqueness or Duplication of Master Data:""")
    cat9.font.size = Pt(12)
    cat9.bold = True
    intt = doc.add_paragraph().add_run("""Duplication of Master Data is a very common area of concern, normally resultant of systems that do not mitigate causes of bad data. Factors involved in capturing Master Data are derived from different procedures; multiple users, locations and languages; data degradation; multiple versions of master data; variation of standards/practices over time and human error""")
    intt = doc.add_paragraph().add_run("""Hence, uniqueness of mater data items is the first check performed on a master data set to determine if processes and systems generate master data that is not duplicated. The lack of unique items or the presence of duplicated items implies that a Governance Structure including master data record management systems solution that forces end users to verify the nonexistence of the item before creation is not in place""")
    intt = doc.add_paragraph().add_run("""Item uniqueness is examined across the following data elements:""")
    intt = doc.add_paragraph().add_run("""\t         i.	Material Numbers""")
    intt = doc.add_paragraph().add_run("""\t        ii.	Descriptions""")
    intt = doc.add_paragraph().add_run("""\t        iii. Manufacturer Names / Supplier Names and Part Numbers""")
    intt = doc.add_paragraph().add_run("""Uniqueness also helps determine the potential initial data set to be used for cataloguing. If material numbers are not determined to be unique, then descriptions will be used. If descriptions are not determined to be unique then a combination of descriptions and manufacturing names / supplier names and part numbers will be used""")
    cat10 = doc.add_paragraph().add_run("""H. Structure of Master Data:""")
    cat10.font.size = Pt(12)
    cat10.bold = True
    intt = doc.add_paragraph().add_run("""Un-Structured data is data that has no relevance. Structuring begins with the development of data requirements for a business function and its related objectives. Master data usage and relevance is reliant upon master data having structure to ensure consistent classification and identification of data to continually eliminate data reconciliation issues. Further, structuring data, complimented with a record Management tool, will contribute to the elimination of item duplication""")
    intt = doc.add_paragraph().add_run("""Lack of structure implies that the business functions within the organization are underperforming in relation to their objectives. Master Data must have uniform structure to ensure the elimination of duplicates or uniqueness, accuracy and relevance. Hence, PiLog will analyze the current data structure used to determine if a deliberate data structure is used to support business objectives""")
    cat11 = doc.add_paragraph().add_run("""I. Standardization of Master Data:""")
    cat11.font.size = Pt(12)
    cat11.bold = True
    intt = doc.add_paragraph().add_run("""Standardized data is reliable and has consistency across multiple items. Data Standards, complemented with Data Structure and uniqueness, provide assurance that data viewed and used is all the data that exists within the master data set, therefore giving the business functions assurance that their business decisions are based on accurate information There are a number of industry standards, such as ISO 8000 that can govern master data and the use of a dataset. This analysis will identify the use of consistent standard rules. Standardization also contributes to elimination of item duplication""")
    cat12 = doc.add_paragraph().add_run("""J. Data Richness:""")
    cat12.font.size = Pt(12)
    cat12.bold = True
    intt = doc.add_paragraph().add_run("""Beyond uniqueness, structure and standard, there is a remaining data principal that indicates the overall quality of master data, Data Richness. Data Richness is the amount of properties available to describe an item and its' related relevance. Relevant properties ensure that master data is traceable, verifiable and complete as it is used by the organization. These relevant properties are described as Mandatory Property Values, which ensure the uniqueness and correct usage of that item, in particular to the purchasing of the item""")
    intt = doc.add_paragraph().add_run("""Often organizational data does not contain all the necessary Mandatory Property Values and this leads to the existence of multiple versions or item duplication. By enriching Master Data property values, especially for critical items, duplications are further reduced; uniqueness is improved; and relevance and completeness of items are increased for the purchasing and inventory management cycles""")
    intt = doc.add_paragraph().add_run("""Our Data Assessment is a services engagement backed by our proprietary algorithms that delivers report findings identifying specific data challenges that may be hindering your operational efficiency and ability to achieve successful business outcomes based on the health of your data.""")
    intt = doc.add_paragraph().add_run("""The ISO 8000 standards are high level requirements that do not prescribe any specific syntax or semantics. On the other hand, the ISO 22745 standards are for a specific implementation of the ISO 8000 standards in extensible mark-up language (XML) and are aimed primarily at parts cataloguing and industrial suppliers""")
    intt = doc.add_paragraph().add_run("""Data Harmonization processes & methodologies complies to ISO 8000 & ISO 22745 standards""")
    doc.add_picture(r'C:\Users\Administrator\Desktop\SEP\data_quality_standards_3.png', width=Inches(6), height=Inches(3.2))

    doc.add_page_break()
    # 3.CONSIDERATIONS ON DATA ANALYSIS:
    line16 = doc.add_paragraph().add_run('3.CONSIDERATIONS ON DATA ANALYSIS:')
    line16.bold = True
    line16.font.size = Pt(18)
    line000 = doc.add_paragraph().add_run('   3.1 Consideration on Data Analysis for DHA:')
    line000.bold = True
    line000.font.size = Pt(14)
    intt = doc.add_paragraph().add_run('         The following considerations or highlights considered for analyzing the Data Health Assessment: ')

    serial_no = ['1','2','3','4','5','6','7','8','9','10','11',]

    Guidelines = ['As some of the material numbers are repeated, the uniqueness of the materials is based on unique descriptions and Record number',
                  'Description length is analyzed based on number of characters in each Descriptions',
                  'Same or Duplicate short Descriptions with different material number are found by sorting all the Short descriptions with respect to Material Numbers',
                  'Same or Duplicate Long Descriptions are found by sorting all the Long descriptions with respect to Material Numbers',
                  'Inconsistencies in the given data is detected by analyzing the data like Part Number is in different formats as "P/N", "PN" etc.',
                  'Material of construction for the spare are not given in standardized way, it is given in full form (like Stainless Steel) and also in short form (like SS)',
                  'The data was analyzed specifically to find out the inconsistencies in Unit of Measure. Itis also in different formats like "Watt, W" etc.',
                  'Data richness on Reference Data Availability (Part Number, Reference Number, Model Number etc.) is Analyzed and counts are given',
                  'Duplication of Master Data w.r.t to Manufacturer Reference numbers like Part number or Model Numbers are analyzed based on genuineness of Part number with Same Manufacturer etc.',
                  'Duplicate Part Number with different Part Type, In this criteria Same part number is provided with different flag type (Suppl P/N, OEM Part No, Mnfr Part No).',
                  'Analysis on Top 10 Commodities and Top 10 UOM are performed by analyzing the Descriptions '

                 ]

    L6 = list(zip(serial_no,Guidelines))
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    row = table.rows[0].cells
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    pp=row[0].add_paragraph('serial_no')
    pp.bold = True
    pp.alignment=WD_TABLE_ALIGNMENT.CENTER
    p=row[1].add_paragraph('Guidelines')
    p.bold = True
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for serial_no, Guidelines in zip(serial_no,Guidelines):
        row = table.add_row().cells
        table.cell(serial_no.index(serial_no),0).width =457200
        p1=row[0].add_paragraph(str(serial_no))
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER

        table.cell(Guidelines.index(Guidelines),1).width =5029200
        p=row[1].add_paragraph(str(Guidelines))
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        # p1.vertical_alignment = WD_ALIGN_VERTICAL.TOP    # 
    intt = doc.add_paragraph().add_run('A good measure of the quality of Master Data is a Median Grade; when 80% of all the Master Data per Characteristic are equal or above the grade B-')
    # doc.add_page_break()

    # page margins 
    sec_pr = doc.sections[0]._sectPr 
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) # apply border changes to section

    # Document header code
    header = doc.sections[0].header
    footer_para = header.paragraphs[0]
    footer_para.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    table_main = header.add_table(1,3, width=Inches(6))
    texts = date_

    tx_cells = table_main.rows[0].cells
    tb_cell_run = tx_cells[0].add_paragraph().add_run()
    tb_cell_run.add_text("    {}\n     ORGANIZATION".format(texts))
    tb_cell_run.font.size = Pt(12)

    pics = table_main.rows[0].cells
    pic = pics[2]
    run = pic.add_paragraph().add_run()
    run.add_picture(r"C:\Users\Administrator\Desktop\pilog\pi.jpg", width=Inches(1))
    last_paragraph = pic.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document footer code 
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para = "_________________________________________________________________________________________________________________________________________________\n \t Document number: PDQH-ORG-DHA-001 \nCopyright 2021 PiLog. This document contains information proprietary to PiLog, India. It may not be reproduced without written permission. Disclosure of the contents of this document to third parties is strictly prohibited. The information herein is correct at time of publication.\n"
    f=footer.add_paragraph()
    f.alignment = 1
    run20 = f.add_run(footer_para)
    run20.font.size = Pt(8)
    def create_element(name):
        return OxmlElement(name)
    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)
    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    add_page_number(doc.sections[0].footer.paragraphs[1].add_run())

    doc.add_page_break()
    # 4.DATA HEALTH AND QUALITY ANALYSIS RESULTS
    line20 = doc.add_paragraph().add_run('4.DATA HEALTH AND QUALITY ANALYSIS RESULTS')
    line21 = doc.add_paragraph().add_run(' 4.1 Detailed Data Analysis on Data Uniqueness')
    line22 = doc.add_paragraph().add_run('   4.1.1 Analysis on Material Records')
    line20.font.size = Pt(18)
    line21.font.size = Pt(16)
    line22.font.size = Pt(14)
    line20.bold = True
    line22.bold = True
    line21.bold = True
    p = doc.add_paragraph('  Report on Total Material Records uniqueness ')
    # data = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
    # data = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
    # data = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
    # data = data.dropna()
    data= pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
    data = data.fillna('')
    MNO_received=data.shape[0]
    MNO_unique =data['MATERIAL'].nunique()
    MNO_dups = MNO_received-MNO_unique-data["MATERIAL"].isnull().sum()
    fig1 = go.Figure()
    fig1.add_trace(go.Bar(
                       x = ['Total Received'],
                       y = [MNO_received],
                       name = 'Total Number of Records',text=[MNO_received],
                        textposition='outside',
                       marker=dict(
                                    color='rgb(64, 135, 94)'
                                  )
     
                    ))
    fig1.add_trace(go.Bar(
                       x = ['Unique'],
                       y = [MNO_unique],
                       name = 'Unique Material Records',text=[MNO_unique],
                        textposition='outside',
                       marker=dict(
                                    color='rgb(168, 162, 64)'
                                  )
       
                    ))
    fig1.add_trace(go.Bar(
                       x = ['Duplicates'],
                       y = [MNO_dups],
                       name = 'Dupilates Material Records',
                        text=[MNO_dups],
                       textposition='outside',
                       marker=dict(
                                    color='rgb(0, 113, 197)'
                                    )
       
                ))
    fig1.update_traces( )
    fig1.update_layout(
                             title={ 'text': "Report on Material Records",'y':0.98,'x':0.5,'xanchor': 'center','yanchor': 'top'},

                xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                titlefont=dict
                (
                            family = 'Times New Roman',
                            size =30,
                            color = 'RED'

                )

    )
    config = {
         'modeBarButtonsToAdd' : ['drawline', 'drawopenpath', 'drawrect', 'eraseshape'],
        'displaylogo'         : False
    }
    total_records_counts = [MNO_received, MNO_unique, MNO_dups]
    xaxis = ['Total Number of Records', 'Unique Material Records', 'Duplicate Material Records']
    L3 = list(zip(xaxis,total_records_counts))
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.allow_autofit = True
    table.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row = table.rows[0].cells
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)

    p=row[0].add_paragraph('Criteria')
    p.alignment=WD_TABLE_ALIGNMENT.CENTER
    p=row[1].add_paragraph('Count')
    p.alignment=WD_TABLE_ALIGNMENT.CENTER
    for Criteria, Count in L3:
        row = table.add_row().cells
        # row[0].text = str(Criteria)
        p2=row[0].add_paragraph(str(Criteria))
        p2.alignment=WD_TABLE_ALIGNMENT.CENTER
        p=row[1].add_paragraph(str(Count))
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    fig1.write_image(r'figA.jpeg',width=9,height=8)
    doc.add_paragraph('\n')
    doc.add_picture(r'C:figA.jpeg', width=Inches(6), height=Inches(3))

    # 4.1.2 Analysis on Descriptions/Texts with prefix 
    doc.add_page_break()
    line23 = doc.add_paragraph().add_run('4.1.2 Analysis on Descriptions/Texts')
    line23.font.size = Pt(14)
    line23.bold = True
    p = doc.add_paragraph('\tReport on Material Descriptions Uniqueness is as below ')

    

    sample = data.dropna(subset=["LONG_DESCRIPTION"])
    MNO_received = sample.shape[0]
    long_received = sample["LONG_DESCRIPTION"].shape[0]
    pp = sample.copy()
    pp["LD"] = pp["LONG_DESCRIPTION"]#.apply(lambda x: re.sub(r"[^0-9A-Za-z\-/]","", x))
    pp["SD"] = pp["SHORT_DESCRIPTION"]#.apply(lambda x: re.sub(r"[^0-9A-Za-z\-/]","", x))

    ld=pp[pp["LD"].isin(list(pp["LD"].value_counts()[pp["LD"].value_counts()>1].index))]
    ld.sort_values(by=["LD"], inplace=True)
    dupes_long = sample.loc[ld.index]

    sdd=pp[pp["SD"].isin(list(pp["SD"].value_counts()[pp["SD"].value_counts()>1].index))]
    sdd.sort_values(by=["SD"], inplace=True)
    dupes_short = sample.loc[sdd.index]

    sd = ld[ld["SD"].isin(list(ld["SD"].value_counts()[ld["SD"].value_counts()>1].index))]
    lis = list(sd[["SD", "LD"]][sd[["SD", "LD"]].duplicated(keep=False)].index)
    outb=sd.loc[lis]
    outp = sample.loc[lis]

    # Total_Records_Descriptions = sample.shape[0]
    Total_Duplicates_on_Short_Description = dupes_short.shape[0]
    Total_Duplicates_on_Long_Description = dupes_long.shape[0]
    Total_Duplicates_on_Short_and_Long_Description = outp.shape[0]


    short_percentage = (Total_Duplicates_on_Short_Description/MNO_received*100)
    short_per = round(short_percentage,2)

    long_percentage = (Total_Duplicates_on_Long_Description/long_received*100)
    long_per = round(long_percentage,2)

    short_long__percentage = (Total_Duplicates_on_Short_and_Long_Description/MNO_received*100)
    short_long__per = round(short_long__percentage,2)


    index=["Duplicates on Short Description","Duplicates on Long Description","Duplicates on Short and Long Description"]
    data_ = {'Count': [Total_Duplicates_on_Short_Description,Total_Duplicates_on_Long_Description,Total_Duplicates_on_Short_and_Long_Description],
    'Percentage':[short_per,long_per,short_long__per]}
    LEN=pd.DataFrame(data=data_,index=index)
    # LEN.drop(LEN[LEN['Count'] ==0].index, inplace = True)
    LEN.reset_index(inplace = True)
    LEN.rename(columns = {"index": "Criteria"}, inplace = True)
    fig22 = px.pie(LEN, values='Count',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    fig22.update_traces(textposition='inside',textinfo='value')
    fig22.update_layout(  title={'text':'Report on Descriptions','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                        titlefont=dict( family='Times New Roman',
                                        size=30,
                                        color='Red'
                                        ),
                        xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                        uniformtext_minsize=20, uniformtext_mode='hide')
    config={'displaylogo':False}
    table = doc.add_table(LEN.shape[0]+1, LEN.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[2]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
    # lst = [4846320,697280]
    for j in range(LEN.shape[-1]):
        # table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(LEN.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    for i in range(LEN.shape[0]):
        row = table.rows[i+1].cells
        for j in range(LEN.shape[-1]):
            # table.cell(i,j).width = lst[j]
            # row = table.rows[j].cells
            p=row[j].add_paragraph(str(LEN.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
    fig22.write_image(r'fig_ref.jpeg')
    doc.add_paragraph('\n')
    doc.add_picture(r'fig_ref.jpeg',width=Inches(6), height=Inches(3))


    # 4.1.3 Advanced Analysis on Descriptions/Texts without prefixes 
    doc.add_page_break()
    sample = data.dropna(subset=["LONG_DESCRIPTION"])
    pp = sample.copy()
    line001 = doc.add_paragraph().add_run('4.1.3 Advanced Analysis on Descriptions/Texts')
    line001.font.size = Pt(14)
    line001.bold = True
    intt = doc.add_paragraph().add_run('Below is the report generated after exclusion of unwanted text in the source description (E.g.: Special characters, Prefixes)')
    part_prefix_df = pd.read_excel(r'C:\Users\Administrator\Desktop\V10_ORIG.xls')
    part_prefix_df = part_prefix_df.dropna()
    part_prefix_df = part_prefix_df.drop([37,40,41,47,48,58])
    mpd = part_prefix_df[part_prefix_df["V10_PART_PREFIX"] != "MANUFACTURER"]
    mpd["V10_PART_PREFIX"].astype('str')
    vten = list(mpd["V10_PART_PREFIX"].astype("str"))
    vten.sort(key = len, reverse=True)
    orig = list(mpd["ORIGINAL_PART_PREFIX"].astype("str"))
    orig.sort(key = len, reverse=True)
    vten.extend(orig)
    pref = list(set(vten))
    pref.sort(key = len, reverse=True)
    pat = r'\b(?:{})\b'.format('|'.join(pref))
    pp["LDWP"]=pp['LONG_DESCRIPTION'].str.replace(pat, '')
    pp["SDWP"]=pp['SHORT_DESCRIPTION'].str.replace(pat, '')
    pp["ld"] = pp["LDWP"].apply(lambda x: re.sub(r"[^0-9A-Za-z\-/]","", x))
    pp["Sd"] = pp["SDWP"].apply(lambda x: re.sub(r"[^0-9A-Za-z\-/]","", x))

    ldwp=pp[pp["ld"].isin(list(pp["ld"].value_counts()[pp["ld"].value_counts()>1].index))]
    ldwp.sort_values(by=["ld"], inplace=True)
    dupes_long_without_prefix = sample.loc[ldwp.index]

    sdwppw=pp[pp["Sd"].isin(list(pp["Sd"].value_counts()[pp["Sd"].value_counts()>1].index))]
    sdwppw.sort_values(by=["Sd"], inplace=True)
    dupes_short_without_prefix = sample.loc[sdwppw.index]

    sdwp = ldwp[ldwp["Sd"].isin(list(ldwp["Sd"].value_counts()[ldwp["Sd"].value_counts()>1].index))]
    lis2=list(sdwp[["Sd", "ld"]][sdwp[["Sd", "ld"]].duplicated(keep=False)].index)
    outb2=sdwp.loc[lis2]
    outp2 = sample.loc[lis2]

    # Total_Records_Descriptions = sample.shape[0]
    Total_Duplicates_on_Short_Description_1 = dupes_short_without_prefix.shape[0]
    Total_Duplicates_on_Long_Description_1 = dupes_long_without_prefix.shape[0]
    Total_Duplicates_on_Short_and_Long_Description_1 = outp2.shape[0]

 
    short_percentage_1 = (Total_Duplicates_on_Short_Description_1/MNO_received*100)
    short_perc = round(short_percentage_1,2)

    long_percentage_1 = (Total_Duplicates_on_Long_Description_1/long_received*100)
    long_perc = round(long_percentage_1,2)

    short_long__percentage_1 = (Total_Duplicates_on_Short_and_Long_Description_1/MNO_received*100)
    short_long__perc = round(short_long__percentage_1,2)


    index=["Duplicates on Short Description","Duplicates on Long Description","Duplicates on Short and Long Description"]
    data_ = {'Count': [Total_Duplicates_on_Short_Description_1,Total_Duplicates_on_Long_Description_1,Total_Duplicates_on_Short_and_Long_Description_1],
    'Percentage':[short_perc,long_perc,short_long__perc]}
    LEN=pd.DataFrame(data=data_,index=index)
    # LEN.drop(LEN[LEN['Count'] ==0].index, inplace = True)
    LEN.reset_index(inplace = True)
    LEN.rename(columns = {"index": "Criteria"}, inplace = True)
    fig223 = px.pie(LEN, values='Count',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    fig223.update_traces(textposition='inside',textinfo='value')
    fig223.update_layout(  title={'text':'Report on Descriptions','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                        titlefont=dict( family='Times New Roman',
                                        size=30,
                                        color='Red'
                                        ),
                        xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                        uniformtext_minsize=20, uniformtext_mode='hide')
    config={'displaylogo':False}
    table = doc.add_table(LEN.shape[0]+1, LEN.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[2]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
    # lst = [4846320,697280]
    for j in range(LEN.shape[-1]):
        # table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(LEN.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    for i in range(LEN.shape[0]):
        row = table.rows[i+1].cells
        for j in range(LEN.shape[-1]):
            # table.cell(i,j).width = lst[j]
            # row = table.rows[j].cells
            p=row[j].add_paragraph(str(LEN.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
    fig223.write_image(r'fig_ref2.jpeg')
    doc.add_paragraph('\n')
    doc.add_picture(r'fig_ref2.jpeg',width=Inches(6), height=Inches(3))


    # 4.1.4 Analysis on Reference Details 
    doc.add_page_break()
    line50 = doc.add_paragraph().add_run('4.1.4 Analysis on Reference Details')
    line50.bold = True
    line50.font.size = Pt(14)
    intt = doc.add_paragraph().add_run('Report on Reference Data Uniqueness is as below')
    # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
    # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
    # sample_df = sample_df[:1500]
    # sample_df = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
    sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
    sample_df = sample_df.drop(columns=["UOM"])
    sample_df = sample_df.dropna(subset=["LONG_DESCRIPTION"])
    sample_df['LONG_DESCRIPTION'] = sample_df['LONG_DESCRIPTION'].apply(str)
    part_prefix_df = pd.read_excel(r'C:\Users\Administrator\Desktop\V10_ORIG.xls')
    # part_prefix_df = part_prefix_df.dropna()
    part_prefix_df = part_prefix_df.fillna('')
    
    

    def potential_duplicates_wrt_manufacturer_part_number_1():
        def part_no_checker(description):
            container = []
            manufacture_part_df = part_prefix_df
            for i, row in manufacture_part_df.iterrows():
                cleaned_part_code = re.sub(":", "", row["ORIGINAL_PART_PREFIX"])
                cleaned_part_code = re.sub("\.", "\\.", cleaned_part_code)
                # pat = re.compile(r"(\b[,;\s]?%s[:\-\s]{0,9}[^,;\s/]*)" % cleaned_part_code)
                pat = re.compile(r"(\b[,;\s]?%s[:\-\s,]{1,9}[a-zA-Z0-9]{1,40}[^,;\s]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b[,;\s]?%s[:\-\s]{1,9}[a-zA-Z0-9]{1,40}[^,;\s/]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b%s[^:]?:\s?[:;,]{0,9}[^,;:]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b;\s?%s[^:]?:\s?[:;,]{0,9}[^,;:]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b;\s?%s[^:]?:\s?[:;,]{0,9}[^,;:]*)" % cleaned_part_code)
                try:
                    match = re.findall(pat, description)
                    for i in range(len(match)):
                        replace_lst = match[i] + "~" + row["V10_PART_PREFIX"]
                        if replace_lst not in container:
                            container.append(replace_lst)

                    # print(replace_lst)
                except TypeError:
                    continue
            if container:
                return container
            else:
                return None

        sample_df_5 = sample_df.dropna(subset=["LONG_DESCRIPTION"])
        sample_df_5["Captured_part_no_list"] = sample_df_5["LONG_DESCRIPTION"].apply(lambda x: part_no_checker(x))
        sample_df_5 = sample_df_5.dropna(subset=["Captured_part_no_list"])
        sample_df_5_explode = sample_df_5.assign(PartNumber=sample_df_5.Captured_part_no_list).explode('PartNumber')
        sample_df_5_explode.drop(columns=["Captured_part_no_list","SHORT_DESCRIPTION"], inplace=True)
        sample_df_5_explode["Reference Type"] = sample_df_5_explode["PartNumber"].apply(lambda x: x.split("~")[1])
        sample_df_5_explode["Total_Reference_Numbers"] = sample_df_5_explode["PartNumber"].apply(lambda x: x.split("~")[0])
        sample_df_5_explode.drop(columns=["PartNumber"], inplace=True)
        return sample_df_5_explode

    REF = potential_duplicates_wrt_manufacturer_part_number_1()

    #################### removing prefixes from the Total Reference Numbers ###################################

    def prefix_remove(x):      
        source = x
        for i in (range(part_prefix_df.shape[0])):
            des = re.sub(str(part_prefix_df['ORIGINAL_PART_PREFIX'][i]),'',source)
            source = des
        return source
    REF['Total_Reference_Numbers'] = REF['Total_Reference_Numbers'].apply(lambda x:prefix_remove(x))
    # REF.to_excel(r'C:\Users\Administrator\Desktop\NOV\reference_numbers_checking_55.xlsx')

    ############################################################################################################
    ##############################below code for removing alphabets , alpha+special charecters and special charecters #########################
    REF["Total_Reference_Numbers"] = REF["Total_Reference_Numbers"].apply(lambda part_numb :re.sub(r'^[\sa-zA-Z]*$','',part_numb))
    REF["Total_Reference_Numbers"] = REF["Total_Reference_Numbers"].apply(lambda part_numb :re.sub(r'^[:\s,.!@#$%^\-&*a-zA-Z;:\s,.!@#$%^\-&*]*$','',part_numb))
    REF["Total_Reference_Numbers"] = REF["Total_Reference_Numbers"].apply(lambda part_numb :re.sub(r'^\W+$','',part_numb))
    REF = REF[REF.Total_Reference_Numbers != '  ']
    REF = REF[REF.Total_Reference_Numbers != ' ']
    REF = REF[REF.Total_Reference_Numbers != '']
    REF = REF.dropna()
    ##################################################################################################################################################

    sorted_df = REF.sort_values(by=['Total_Reference_Numbers'], ascending=True)
    total_dup = sorted_df[sorted_df["Total_Reference_Numbers"].isin(list(sorted_df["Total_Reference_Numbers"].value_counts()[sorted_df["Total_Reference_Numbers"].value_counts()>1].index))]
    print(total_dup)

    ####### TO UPDATE DATA INTO DATABASE ########

    total_dup["BATCH_ID"] = BATCH_ID

    total_dup["TABLE_TYPE"] = "Reference_Details"
    
    total_dup['INSERT_BY'] = '{}'.format(insert_time)

    total_dup = total_dup.assign(VAL_SEQUENCE=[1 + i for i in range(len(total_dup))])[['VAL_SEQUENCE'] + total_dup.columns.tolist()]


    total_dup = total_dup[['MATERIAL','LONG_DESCRIPTION','Reference Type','Total_Reference_Numbers','BATCH_ID','TABLE_TYPE','VAL_SEQUENCE', 'INSERT_BY']]

    sql = """insert into CLOUD_DHA(MATERIAL,LONGDESCRIPTION,REFERENCE_TYPE,EXTRACTED_REFERENCE_NUMBERS,BATCH_ID,TABLE_TYPE,VAL_SEQUENCE, INSERT_DATE ) values(:1, :2, :3, :4, :5, :6, :7, :8)"""

    cursor = conn.cursor()
    for c, r in total_dup.iterrows():
        cursor.execute(sql, tuple(r))
        conn.commit()

    #############################################

    gru_by = total_dup.groupby(by = 'Total_Reference_Numbers',as_index=False)
    cnt2 = []
    cnt = []
    for i in total_dup['Total_Reference_Numbers'].unique():
        if len(gru_by.get_group(i)['Reference Type'].value_counts())>1:
            cnt.append(sum(gru_by.get_group(i)['Reference Type'].value_counts()))
            cnt2.append(len((gru_by.get_group(i)['Reference Type'].unique())))

    try:
        Total_Duplicates_across_different_Reference_Types = sum(cnt)
    except:
        Total_Duplicates_across_different_Reference_Types = 0

    try:
        part_no = total_dup["Reference Type"].value_counts()['MANUFACTURER PART NO']
    except:
        part_no = 0

    try:
        model_no = total_dup["Reference Type"].value_counts()['MODEL/MACHINE NO']
    except:
        model_no = 0

    try:
        Reference_no = total_dup["Reference Type"].value_counts()['REFERENCE NO']
    except:
        Reference_no = 0


    part_percentage = (part_no/MNO_received*100)
    part_per = round(part_percentage,2)
    # part_per = str(part_per) + "%"

    Reference_percentage = (Reference_no/MNO_received*100)
    Refere_per = round(Reference_percentage,2)
    # Refere_per = str(Refere_per) + "%"

    model__percentage = (model_no/MNO_received*100)
    model__per = round(model__percentage,2)
    # model__per = str(model__per) + "%"

    Different_reference_types = (Total_Duplicates_across_different_Reference_Types/MNO_received*100)
    Different_reference_types_per = round(Different_reference_types,2)
    # Different_reference_types_per = str(Different_reference_types_per) + "%"

    index=["Total Duplicates on Part Numbers","Total Duplicates on Model Numbers","Total Duplicates on Reference Numbers","Total Duplicates across different Reference Types"]
    data_ = {'Count': [part_no,model_no,Reference_no,Total_Duplicates_across_different_Reference_Types],
    'Percentage':[part_per,model__per,Refere_per,Different_reference_types_per]}
    LEN=pd.DataFrame(data=data_,index=index)
    # LEN.drop(LEN[LEN['Count'] ==0].index, inplace = True)
    LEN.reset_index(inplace = True)
    LEN.rename(columns = {"index": "Criteria"}, inplace = True)
    fig223 = px.pie(LEN, values='Count',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    fig223.update_traces(textposition='inside',textinfo='value')
    fig223.update_layout(  title={'text':'Report on Reference Data Completeness','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                        titlefont=dict( family='Times New Roman',
                                        size=30,
                                        color='Red'
                                        ),
                        xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                        uniformtext_minsize=20, uniformtext_mode='hide')
    config={'displaylogo':False}
    table = doc.add_table(LEN.shape[0]+1, LEN.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[2]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
    # lst = [4846320,697280]
    for j in range(LEN.shape[-1]):
        # table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(LEN.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    for i in range(LEN.shape[0]):
        row = table.rows[i+1].cells
        for j in range(LEN.shape[-1]):
            # table.cell(i,j).width = lst[j]
            # row = table.rows[j].cells
            p=row[j].add_paragraph(str(LEN.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
    fig223.write_image(r'fig_ref9.jpeg')
    # except:
    #     doc.add_paragraph('')





    # ref_dups_ans = ["Count", str(part_no), str(model_no), str(Reference_no),str(Total_Duplicates_across_different_Reference_Types)]
    # ref_dups = ['Criteria',
    #  'Total Duplicates on Part Numbers',
    #  'Total Duplicates on Model Numbers',
    #  'Total Duplicates on Reference Numbers',
    #  'Total Duplicates across different Reference Types']

    # # ref_dups_ans = ["Count", str(one), str(two), str(three),str(four)]
    # table = doc.add_table(5,2,style = 'Table Grid')
    # row = table.rows[0].cells[1]
    # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    # table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    # table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)

    # for j in range(len(ref_dups)):
    #     table.cell(j,0).width = 5212080
    # #     table.cell(j,0).text = ref_dups[j]
    #     row = table.rows[j].cells
    #     p=row[0].add_paragraph(ref_dups[j])
    #     p.alignment=WD_TABLE_ALIGNMENT.CENTER
        
    # for j in range(len(ref_dups_ans)):
    #     table.cell(j,1).width = 2103120
    # #     table.cell(j,1).text = ref_dups_ans[j]
    #     row = table.rows[j].cells
    #     p=row[1].add_paragraph(ref_dups_ans[j])
    #     p.alignment=WD_TABLE_ALIGNMENT.CENTER

    # ref_dups = {'count':[ str(part_no), str(model_no), str(Reference_no),str(Total_Duplicates_across_different_Reference_Types)]}
    # index = [
    #  'Total Duplicates on Part Numbers',
    #  'Total Duplicates on Model Numbers',
    #  'Total Duplicates on Reference Numbers',
    #  'Total Duplicates across different Reference Types']

    # LEN=pd.DataFrame(data=ref_dups,index=index)
    # LEN.reset_index(inplace = True)
    # LEN.rename(columns = {"index": "Criteria"}, inplace = True)
    # fig22 = px.pie(LEN, values='count',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    # fig22.update_traces(textposition='inside',textinfo='value')
    # fig22.update_layout(  title={'text':'Report on Reference Details','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
    #                     titlefont=dict( family='Times New Roman',
    #                                     size=30,
    #                                     color='Red'
    #                                     ),
    #                     xaxis=dict(
    #                         showgrid=True,
    #                         showline=True,
    #                         showticklabels=True,
    #                         zeroline=True,
    #                         linecolor='Black',
    #                         linewidth=2,
    #                         gridcolor= 'white',
    #                         zerolinecolor= 'Black'
    #                         #domain=[0.15, 1]
    #                     ),
    #             yaxis=dict(
    #                         showgrid=True,
    #                         showline=True,
    #                         showticklabels=True,
    #                         zeroline=True,
    #                         linecolor='Black',
    #                         linewidth=2,
    #                         gridcolor= 'White',
    #                         zerolinecolor= 'Black'
    #                     ),
    #                         showlegend=True,
    #                         paper_bgcolor='rgb(248, 248, 255)',
    #                         plot_bgcolor='rgb(248, 246, 255)',
    #                         margin=dict(l=100, r=20, t=70, b=70),
    #                     uniformtext_minsize=20, uniformtext_mode='hide')
    # config={'displaylogo':False}
    # fig22.write_image(r'C:\Users\Administrator\Desktop\NOV/fig333.jpeg')

    ############## for adding hyperlink below the table ###############
    def add_hyperlink(paragraph, url, text, color, underline):

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if not color is None:
          c = docx.oxml.shared.OxmlElement('w:color')
          c.set(docx.oxml.shared.qn('w:val'), color)
          rPr.append(c)

        # Remove underlining if it is requested
        if not underline:
          u = docx.oxml.shared.OxmlElement('w:u')
          u.set(docx.oxml.shared.qn('w:val'), 'none')
          rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink
    p = doc.add_paragraph("To download complete Reference Details Data ")
    hyperlink = add_hyperlink(p, 'http://apihub.piloggroup.com:6652/user_down/Reference_Details/{}/{}/Reference Details {}'.format(BATCH_ID, insert_time, format(datetime.now().replace(microsecond=0))), 'click here', '0000FF', True)

    ###########################################################################################
    doc.add_paragraph('\n')
    doc.add_picture(r'fig_ref9.jpeg',width=Inches(6), height=Inches(3))

    now = datetime.now()
    print(' 4.1.5 start time ',now)

    #4.1.5 Potential Duplicates w.r.t Manufacturer Part Number
    line32 = doc.add_paragraph().add_run('4.1.5 Potential Duplicates w.r.t Reference Number')
    line32.bold = True
    line32.font.size = Pt(14)

    # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
    # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
    sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
    # sample_df = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
    sample_df = sample_df.drop(columns=["UOM","SHORT_DESCRIPTION"])
    sample_df = sample_df.dropna(subset=["LONG_DESCRIPTION"])
    # sample_df = np.array_split(sample_df,3)
    rgb_list = ["E6E6FA", "F0FFF0", "FFF5EE", "EEE8AA", "87CEEB", "AFEEEE", "E0FFFF","ADD8E6","87CEEB","5F9EA0","B0E0E6","AFEEEE","40E0D0","66CDAA","F08080","FFA07A","EE82EE","FFE4C4","FFFFE0","DCDCDC","F8F8FF"]

    part_prefix_df = pd.read_excel(r'C:\Users\Administrator\Desktop\V10_ORIG.xls')
    part_prefix_df = part_prefix_df.fillna('')
    part_prefix_df['length'] = part_prefix_df['ORIGINAL_PART_PREFIX'].str.len()

    part_prefix_df.sort_values('length', ascending=False, inplace=True)

    part_prefix_df = part_prefix_df.reset_index(drop = True)



    def potential_duplicates_wrt_manufacturer_part_number_1():
        def part_no_checker(description):
            container = []
            manufacture_part_df = part_prefix_df[part_prefix_df["V10_PART_PREFIX"] != "MANUFACTURER"]
            for i, row in manufacture_part_df.iterrows():
                cleaned_part_code = re.sub(":", "", row["ORIGINAL_PART_PREFIX"])
                cleaned_part_code = re.sub("\.", "\\.", cleaned_part_code)
                pat = re.compile(r"(\b[,;\s]?%s[:\-\s,]{1,9}[a-zA-Z0-9]{1,40}[^,;\s]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b%s[^:]?:\s?[:;,]{0,9}[^,;:]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b;\s?%s[^:]?:\s?[:;,]{0,9}[^,;:]*)" % cleaned_part_code)
                # pat = re.compile(r"(\b%s.?[:]\s?:?:?:?,?,?,?[^,:;]*[:,;])" % cleaned_part_code)
                # pat = re.compile(r"(\b%s.?[:][^,:]*[:,])" % cleaned_part_code)
                try:
                    match = re.search(pat, description)
                    if match:
                        #below line for removing finded item
                        description = re.sub(re.escape(str(re.search(pat,description).group().strip())), '', description)
                except TypeError:
                    continue
                if match:
                    captured_match = re.sub(r"[^a-zA-Z0-9:]", "", match.group().split("~")[0])
                    captured_match = match.group().split(":")[0] + "~" + captured_match
                    # captured_match = re.sub(r"[^a-zA-Z0-9:]", "", match.group().split(":")[0])
                    # captured_match = match.group() + "~" + row["V10_PART_PREFIX"]
                    if captured_match not in container:
                        container.append(captured_match)
            if container:
                return container
            else:
                return None

        sample_df_5 = sample_df.dropna(subset=["LONG_DESCRIPTION"])
        sample_df_5["Captured_part_no_list"] = sample_df_5["LONG_DESCRIPTION"].apply(lambda x: part_no_checker(x))

        sample_df_5 = sample_df_5.dropna(subset=["Captured_part_no_list"])

        sample_df_5_explode = sample_df_5.assign(ManufacturerPartNumber=sample_df_5.Captured_part_no_list).explode(
            'ManufacturerPartNumber')

        sample_df_5_explode.drop(columns=["Captured_part_no_list"], inplace=True)

        sample_df_5_explode["trimmed_part"] = sample_df_5_explode["ManufacturerPartNumber"].apply(lambda x: x.split("~")[1])


        sample_df_5_explode["ReferenceNumber"] = sample_df_5_explode["ManufacturerPartNumber"].apply(lambda x: x.split("~")[0])
        sample_df_5_output = sample_df_5_explode[sample_df_5_explode.duplicated("trimmed_part", keep=False)]
        sample_df_5_output = sample_df_5_output.drop(columns=["trimmed_part","ManufacturerPartNumber"])
        sample_df_5_output = sample_df_5_output.sort_values(by=["ReferenceNumber"])
        return sample_df_5_output

    sample2 = potential_duplicates_wrt_manufacturer_part_number_1()
    # print(sample2)
    # print("1111111111111111")
    ######################################################################################################
    def prefix_remove(x):      
        source = x
        for i in (range(part_prefix_df.shape[0])):
            des = re.sub(str(part_prefix_df['ORIGINAL_PART_PREFIX'][i]),'',source)
            source = des
        return source
    sample2['ReferenceNumber'] = sample2['ReferenceNumber'].apply(lambda x:prefix_remove(x))
    
    #########################################################################################################
    ##############################below code for removing alphabets , alpha+special charecters and special charecters #########################
    sample2["ReferenceNumber"] = sample2["ReferenceNumber"].apply(lambda part_numb :re.sub(r'^[\sa-zA-Z]*$','',part_numb))
    sample2["ReferenceNumber"] = sample2["ReferenceNumber"].apply(lambda part_numb :re.sub(r'^[:\s,.!@#$%^\-&*a-zA-Z;:\s,.!@#$%^\-&*]*$','',part_numb))
    sample2["ReferenceNumber"] = sample2["ReferenceNumber"].apply(lambda part_numb :re.sub(r'^\W+$','',part_numb))
    sample2 = sample2[sample2.ReferenceNumber != '  ']
    sample2 = sample2[sample2.ReferenceNumber != ' ']
    sample2 = sample2[sample2.ReferenceNumber != '']
    sample2 = sample2.dropna()
    ##################################################################################################################################################



    # sample2["ManufacturerPartNumber"] = sample2["ManufacturerPartNumber"].apply(lambda part_numb :re.sub(r'\b[\sa-zA-Z]*\b','',part_numb))
    # sample2 = sample2[sample2.ManufacturerPartNumber != '  ']

    # sample2['ManufacturerPartNumber'] = sample2['ManufacturerPartNumber'].apply(lambda x: x[0:-1])
    unique_part_no1 = sample2["ReferenceNumber"].unique().tolist()
    sample2.rename(columns={'ReferenceNumber':'Reference_Number'},inplace=True)
    sample2.sort_values(by=["Reference_Number"], inplace=True,ascending=False)

    sample2["Reference_Number"] = sample2.Reference_Number.str.replace(r'\b([\w\W]{1,4})\b', '')

    sample2["Reference_Number"].fillna('', inplace = True)

    sample2 = sample2[sample2['Reference_Number'] != '  ']

    sac=sample2.copy()
    sac.reset_index(inplace=True)

    cnt=0
    datasize = 0
    for i in range(len(sac)):
        try:
            if sac["Reference_Number"][i] != sac["Reference_Number"][1+i]:
                cnt=cnt+1
                if cnt==5:
                    datasize=i
                    break
                    # if cnt==5:
                    #     break
                else:
                    continue
            # if cnt==5:
            #     break
        except KeyError:
            continue

    sample2 = sample2[0:datasize-2]

    print('!!!!!!!!!!!!!!!!!!!!',sample2)
    rgb_dict = dict(zip(unique_part_no1, rgb_list))

    table = doc.add_table(sample2.shape[0]+1, sample2.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[2]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

    # lst = [914400,3657600,914400]

    for j in range(sample2.shape[-1]):
        # table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(sample2.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER
        # table.cell(0,j).text = sample5.columns[j]

    for i in range(sample2.shape[0]):
        row = table.rows[i+1].cells
        for j in range(sample2.shape[-1]):
            # table.cell(i+1,j).width = lst[j
            p=row[j].add_paragraph(str(sample2.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
            # table.cell(i+1,j).text = str(sample5.values[i,j])
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), rgb_dict.get(sample2.values[i,2])))
            table.cell(i+1,j)._tc.get_or_add_tcPr().append(shading_elm_2)
    # except:
    #     doc.add_paragraph('')

    doc.add_paragraph('\n')
    now = datetime.now()
    print(' 4.1.5 end time ',now)
    doc.add_page_break()
    # 4.2 Detailed Data Analysis on Data Completeness
    line003 = doc.add_paragraph().add_run('4.2 Detailed Data Analysis on Data Completeness ')
    line003.font.size = Pt(14)
    line003.bold = True
    line004 = doc.add_paragraph().add_run('4.2.1 Analysis on Description Length ')
    line004.font.size = Pt(14)
    line004.bold = True
    intt = doc.add_paragraph().add_run('Below is the Analysis on description lengths of short descriptions ')

    length=data['SHORT_DESCRIPTION']
    length1=pd.DataFrame(length)
    length1['Count'] = length1['SHORT_DESCRIPTION'].str.len()
    length1 = length1.fillna(0)
    length1['Count']=length1['Count'].astype(int)

    SD_Length_15 = np.sum(length1['Count']<15)
    SD_Length_15_30 = np.sum(length1['Count'].isin(range(16,30)))
    SD_Length_30_40 = np.sum(length1['Count'].isin(range(31,40)))

    index=["Materials having Short text length less than 10","Materials having Short text length between 10 and 30","Materials having Short text length between 30 and 40"]
    data = {'Count': [SD_Length_15,SD_Length_15_30,SD_Length_30_40]}

    LEN=pd.DataFrame(data=data,index=index)
    LEN.drop(LEN[LEN['Count'] ==0].index, inplace = True)
    LEN.reset_index(inplace = True)
    LEN.rename(columns = {"index": "Criteria"}, inplace = True)

    fig22 = px.pie(LEN, values='Count',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    fig22.update_traces(textposition='inside',textinfo='value')
    fig22.update_layout(  title={'text':'Report on Short Description Length','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                        titlefont=dict( family='Times New Roman',
                                        size=30,
                                        color='Red'
                                        ),
                        xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                        uniformtext_minsize=20, uniformtext_mode='hide')
    config={'displaylogo':False}
    fig22.write_image(r'fig36.jpeg')

    table = doc.add_table(LEN.shape[0]+1, LEN.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[1]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)

    lst = [4846320,697280]

    for j in range(LEN.shape[-1]):
        table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(LEN.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    for i in range(LEN.shape[0]):
        row = table.rows[i+1].cells
        for j in range(LEN.shape[-1]):
            table.cell(i,j).width = lst[j]
            # row = table.rows[j].cells
            p=row[j].add_paragraph(str(LEN.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER

    doc.add_paragraph('\n')

    # for min, max,medium count code
    line0004 = doc.add_paragraph().add_run('Analysis on Min,Max,Medium Lengths ')
    line0004.font.size = Pt(14)
    line0004.bold = True
 
    one = int(SD_Length_15)
    two = int(SD_Length_15_30)
    three = int(SD_Length_30_40)
    dt = pd.DataFrame([one, two, three], columns=['new_one'])

    min_value_column = int(dt['new_one'].min())
    medum_value_column = int(dt["new_one"].median())
    max_value_column = int(dt['new_one'].max())

    index=["Minimum short text length","Medium short text length","Maximum Short text length"]
    data = {'Length': [min_value_column,medum_value_column,max_value_column]}

    LEN = pd.DataFrame(data=data,index=index)
    LEN.drop(LEN[LEN['Length'] ==0].index, inplace = True)
    LEN.reset_index(inplace = True)
    LEN.rename(columns = {"index": "Criteria"}, inplace = True)

    table = doc.add_table(LEN.shape[0]+1, LEN.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[1]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)

    lst = [4846320,697280]
    for j in range(LEN.shape[-1]):
        table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(LEN.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    for i in range(LEN.shape[0]):
        row = table.rows[i+1].cells
        for j in range(LEN.shape[-1]):
            table.cell(i,j).width = lst[j]
            # row = table.rows[j].cells
            p=row[j].add_paragraph(str(LEN.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER

    doc.add_paragraph('\n')
    doc.add_picture(r'fig36.jpeg',width=Inches(6), height=Inches(3))

    # Analysis on Description Length (PERCENTAGE WISE)
    intt = doc.add_paragraph().add_run('Below is the Analysis on description lengths of  long descriptions ')
    # data = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
    # data = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
    # data = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
    data = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
    data = data.dropna(subset=['LONG_DESCRIPTION'])
    lengths_long = data["LONG_DESCRIPTION"].apply(lambda x: len(x))
    maxi = max(lengths_long)
    first_group = round(maxi*0.25)-round(maxi*0.25)%10
    second_group = 2*first_group
    third_group = 3*first_group
    first_count=lengths_long.between(0, first_group, inclusive="right")
    second_count=lengths_long.between(first_group+1, second_group, inclusive="right")
    third_count=lengths_long.between(second_group+1, third_group, inclusive="right")
    fourth_count=lengths_long.between(third_group+1, maxi, inclusive="right")

    
    phenom_ld_len_data = pd.DataFrame({"Criteria":["Long Descriptions with character length ranges between 0 and {}".format(first_group), "Long Descriptions with character length ranges between {} and {}".format(first_group, second_group), "Long Descriptions with character length ranges between {} and {}".format(second_group, third_group), "Long Descriptions with character length ranges above {}".format(third_group)]
    , "Count":[first_count.value_counts()[1],second_count.value_counts()[1], third_count.value_counts()[1],fourth_count.value_counts()[1]]})

    phenom_ld_len_data_graph = pd.DataFrame({"Criteria":["0 to {}".format(first_group), "{} to {}".format(first_group, second_group), "{} to {}".format(second_group, third_group), "above {}".format(third_group)]
    , "Count":[first_count.value_counts()[1],second_count.value_counts()[1], third_count.value_counts()[1],fourth_count.value_counts()[1]]})


    phenom_ld_len_data.drop(phenom_ld_len_data[phenom_ld_len_data['Count'] ==0].index, inplace = True)
    fig222 = px.pie(phenom_ld_len_data_graph, values='Count',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    fig222.update_traces(textposition='inside',textinfo='value')
    fig222.update_layout(  title={'text':'Report on Long Descriptions','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                        titlefont=dict( family='Times New Roman',
                                        size=30,
                                        color='Red'
                                        ),
                        xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                        uniformtext_minsize=20, uniformtext_mode='hide')
    config={'displaylogo':False}
    table = doc.add_table(phenom_ld_len_data.shape[0]+1, phenom_ld_len_data.shape[1],style = 'Table Grid')
    row=table.rows[0].cells[1]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    # table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
    # lst = [4846320,697280]
    for j in range(phenom_ld_len_data.shape[-1]):
        # table.cell(0,j).width = lst[j]
        col = table.columns[j].cells
        p=col[0].add_paragraph(phenom_ld_len_data.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

    for i in range(phenom_ld_len_data.shape[0]):
        row = table.rows[i+1].cells
        for j in range(phenom_ld_len_data.shape[-1]):
            # table.cell(i,j).width = lst[j]
            # row = table.rows[j].cells
            p=row[j].add_paragraph(str(phenom_ld_len_data.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
    fig222.write_image(r'fig_ref_.jpeg')
    doc.add_paragraph('\n')
    doc.add_picture(r'fig_ref_.jpeg',width=Inches(6), height=Inches(3))



    #4.2.2 report on reference numbers count
    try:
        line006 = doc.add_paragraph().add_run('4.2.2 Report on Reference Data Completeness')
        line006.font.size = Pt(14)
        line006.bold = True
        
        try:
            part_number = REF['Reference Type'].value_counts()['MANUFACTURER PART NO']
        except:
            part_number = 0
        try:
            MODEL_no = REF['Reference Type'].value_counts()['MODEL/MACHINE NO']
        except:
            MODEL_no = 0

        try:
            REFERENCE_NO = REF['Reference Type'].value_counts()['REFERENCE NO']
        except:
            REFERENCE_NO = 0

        try:
            DRAWING_no = REF['Reference Type'].value_counts()['DRAWING']
        except:
            DRAWING_no = 0

        try:
            vendor = REF['Reference Type'].value_counts()['MANUFACTURER']
        except:
            vendor = 0

        try:
            SUPPLIER = REF['Reference Type'].value_counts()['SUPPLIER PART NO']
        except:
            SUPPLIER = 0

        try:
            OEM_PART_NO = REF['Reference Type'].value_counts()['OEM PART NO']
        except:
            OEM_PART_NO = 0

        Partnumber_per = (part_number/MNO_received*100)
        Partnumber_percentage = round(Partnumber_per,2)

        model_per = (MODEL_no/MNO_received*100)
        model_percentage = round(model_per,2)

        Reference_number_per = (REFERENCE_NO/MNO_received*100)
        Reference_number_percentage = round(Reference_number_per,2)

        Drawing_number_per = (DRAWING_no/MNO_received*100)
        Drawing_number_percentage = round(Drawing_number_per,2)

        Vendor_per = (vendor/MNO_received*100)
        Vendor_percentage = round(Vendor_per,2)

        Supplier_per = (SUPPLIER/MNO_received*100)
        Supplier_percentage = round(Supplier_per,2)

        OEM_Part_per = (OEM_PART_NO/MNO_received*100)
        OEM_Part_percentage = round(OEM_Part_per,2)

        index=["Total part numbers","Total model numbers","Toatal reference numbers","Total drawing numbers","Total vendor names","Total supplier","Total OEM Part numbers"]
        data = {'Total Number of Records': [part_number,MODEL_no,REFERENCE_NO,DRAWING_no,vendor,SUPPLIER,OEM_PART_NO],
        'Percentage':[Partnumber_percentage,model_percentage,Reference_number_percentage,Drawing_number_percentage,Vendor_percentage,Supplier_percentage,OEM_Part_percentage]}
        LEN=pd.DataFrame(data=data,index=index)
        LEN.drop(LEN[LEN['Total Number of Records'] ==0].index, inplace = True)
        LEN.reset_index(inplace = True)
        LEN.rename(columns = {"index": "Criteria"}, inplace = True)
        fig22 = px.pie(LEN, values='Total Number of Records',names='Criteria',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
        fig22.update_traces(textposition='inside',textinfo='value')
        fig22.update_layout(  title={'text':'Report on Reference Data Completeness','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                            titlefont=dict( family='Times New Roman',
                                            size=30,
                                            color='Red'
                                            ),
                            xaxis=dict(
                                showgrid=True,
                                showline=True,
                                showticklabels=True,
                                zeroline=True,
                                linecolor='Black',
                                linewidth=2,
                                gridcolor= 'white',
                                zerolinecolor= 'Black'
                                #domain=[0.15, 1]
                            ),
                    yaxis=dict(
                                showgrid=True,
                                showline=True,
                                showticklabels=True,
                                zeroline=True,
                                linecolor='Black',
                                linewidth=2,
                                gridcolor= 'White',
                                zerolinecolor= 'Black'
                            ),
                                showlegend=True,
                                paper_bgcolor='rgb(248, 248, 255)',
                                plot_bgcolor='rgb(248, 246, 255)',
                                margin=dict(l=100, r=20, t=70, b=70),
                            uniformtext_minsize=20, uniformtext_mode='hide')
        config={'displaylogo':False}
        table = doc.add_table(LEN.shape[0]+1, LEN.shape[1],style = 'Table Grid')
        row=table.rows[0].cells[2]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
        table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
        table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
        # lst = [4846320,697280]
        for j in range(LEN.shape[-1]):
            # table.cell(0,j).width = lst[j]
            col = table.columns[j].cells
            p=col[0].add_paragraph(LEN.columns[j])
            p.alignment=WD_TABLE_ALIGNMENT.CENTER

        for i in range(LEN.shape[0]):
            row = table.rows[i+1].cells
            for j in range(LEN.shape[-1]):
                # table.cell(i,j).width = lst[j]
                # row = table.rows[j].cells
                p=row[j].add_paragraph(str(LEN.values[i,j]))
                p.alignment=WD_TABLE_ALIGNMENT.CENTER
        fig22.write_image(r'fig_ref_count.jpeg')
        doc.add_paragraph('\n')
        doc.add_picture(r'fig_ref_count.jpeg',width=Inches(6), height=Inches(3))
    except:
        doc.add_paragraph('')

    doc.add_page_break()
    #4.3 Detailed Data Analysis on Data Consistency
    try:
        # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
        # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
        sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
        # sample_df = sample_df[:1500]
        # sample_df = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
        sample_df = sample_df.drop(columns=["UOM","SHORT_DESCRIPTION"])
        sample_df = sample_df.dropna(subset=["LONG_DESCRIPTION"])
        line006 = doc.add_paragraph().add_run('4.3 Detailed Data Analysis on Data Consistency')
        line006.font.size = Pt(14)
        line006.bold = True
        rgb_list = ["E6E6FA", "F0FFF0", "FFF5EE", "EEE8AA", "7FFFD4", "696969", "E0FFFF","ADD8E6","87CEEB","5F9EA0","B0E0E6","AFEEEE","40E0D0","66CDAA","F08080","FFA07A","EE82EE","FFE4C4","FFFFE0","DCDCDC","F8F8FF"]
        def inconsistencies_in_the_description():
            def consistency_checker(description):
                container = []
                for i, row in part_prefix_df.iterrows():
                    if row.V10_PART_PREFIX.lower() != "manufacturer":
                        cleaned_part_code = re.sub(":", "", row["ORIGINAL_PART_PREFIX"])
                        cleaned_part_code = re.sub("\.", "\\.", cleaned_part_code)
                        pat = re.compile(r"(\b%s.?[:].....[^, ]*)" % cleaned_part_code)
                        try:
                            match = re.search(pat, description)
                        except TypeError:
                            continue
                        if match:
                            captured_match = row.V10_PART_PREFIX + ":" + match.group()
                            if captured_match not in container:
                                container.append(captured_match)
                if container:
                    return container
                else:
                    return None
                
            sample_df_2 =sample_df["LONG_DESCRIPTION"].apply(lambda x: re.sub(r" :",":",x))
            sample_2_df = pd.DataFrame(sample_df_2)
            sample_2_df = sample_2_df.drop_duplicates(subset=["LONG_DESCRIPTION"])
            sample_2_df["Captured_part_list"] = sample_2_df["LONG_DESCRIPTION"].apply(lambda x: consistency_checker(x))
            sample_2_df = sample_2_df.dropna(subset=["Captured_part_list"])
            sample_2_df_explode = sample_2_df.assign(Captured_parts=sample_2_df.Captured_part_list).explode('Captured_parts')
            sample_2_df_explode["Captured_V010"] = sample_2_df_explode["Captured_parts"].apply(lambda x: x.split(":")[0])
            sample_2_df_explode["Variant Format"] = sample_2_df_explode["Captured_parts"].apply(lambda x: x.split(":")[1])

            ####### for Number of Materials Linked column #########
            samp = sample_2_df_explode["Variant Format"]
            sam_df = pd.DataFrame(samp)
            sam_df['MATERIAL']=sample_df['MATERIAL']
            def mat_Count(i):
                return sam_df["Variant Format"].value_counts()[i]
            sam_df['Number of Materials Linked']=sam_df["Variant Format"].apply(lambda x:mat_Count(x))
            sam_df.drop(columns=["Variant Format"], inplace=True)
            ###########################################################

            sample_2_df_explode["Captured_Post"] = sample_2_df_explode["Captured_parts"].apply(lambda x: x.split(":")[2])
            sample_2_df_explode.drop(columns=["Captured_part_list", "Captured_parts", "Captured_Post"], inplace=True)
            # sample_2_df_explode.drop_duplicates(subset=["Variant Format"],inplace=True)
            ####################
            frames = [sam_df,sample_2_df_explode]
            result = pd.concat(frames,axis=1)
            ##################

            list_to_concat_df = []
            for item in result["Captured_V010"].unique().tolist():
                if len(result[result["Captured_V010"] == item]["Variant Format"].unique().tolist()) > 1:
                    list_to_concat_df.append(result[result["Captured_V010"] == item])

            if list_to_concat_df:
                sample_2_df_output = pd.concat(list_to_concat_df)
                
                sample_2_df_output["Standard Format"] = sample_2_df_output.apply(
                    lambda x: "%s" % (x["Captured_V010"]), axis=1)
                sample_2_df_output.drop_duplicates(subset=["Variant Format"],inplace=True)
                sample_2_df_output.drop(columns=["Captured_V010"], inplace=True)
                sample_2_df_output.drop(columns=["LONG_DESCRIPTION"], inplace=True)
                sample_2_df_output.sort_values(by=["Variant Format"], inplace=True)
                sample_2_df_output = sample_2_df_output.reset_index()
                sample_2_df_output = sample_2_df_output.rename(columns={"index":"Serial no"})
                sample_2_df_output['Serial no'] = sample_2_df_output.index + 1

                return sample_2_df_output
            else:
                return None
        sample5 = inconsistencies_in_the_description()
        
        # ###################### below code is for previous no of materials linked column ############################
        # sample55 = sample5.copy()

        # def UOM_Count(i):
        #     return sample5["Standard Format"].value_counts()[i]

        # sample5['Number of Materials Linked']=sample5["Standard Format"].apply(lambda x:UOM_Count(x))

        # def UOM_Counts(i):
        #     return sample55["Standard Format"].value_counts()[i]

        # sample55['Number of Materials Linked']=sample55["Standard Format"].apply(lambda x:UOM_Counts(x))
        # ##########################################################################################################################

        # #######  TO UPDATE DATA INTO DATABASE  ##########
        sample55 = sample5.copy()
        sample55["BATCH_ID"] = BATCH_ID

        sample55["TABLE_TYPE"] = "inconsistencies_Prefixes"
        
        sample55['INSERT_BY'] = '{}'.format(insert_time)

        sample55 = sample55.assign(VAL_SEQUENCE=[1 + i for i in range(len(sample55))])[['VAL_SEQUENCE'] + sample55.columns.tolist()]

        sample55 = sample55[['Standard Format','Variant Format','Number of Materials Linked','BATCH_ID','TABLE_TYPE','VAL_SEQUENCE', 'INSERT_BY']]

        sql = """insert into CLOUD_DHA(STANDARD_FORMAT,VARIANT_FORMAT,NUMBER_OF_MATERIALS_LINKED,BATCH_ID,TABLE_TYPE,VAL_SEQUENCE, INSERT_DATE ) values(:1, :2, :3, :4, :5, :6,:7)"""

        cursor = conn.cursor()
        for c, r in sample55.iterrows():
            cursor.execute(sql, tuple(r))
            conn.commit()

        # #############################################################
        sample5 = sample5[['Standard Format','Variant Format','Number of Materials Linked']]
        unique_desc = sample5["Standard Format"].unique().tolist()
        rgb_dict = dict(zip(unique_desc, rgb_list))
        table = doc.add_table(sample5.shape[0]+1,sample5.shape[1],style = 'Table Grid')
        row = table.rows[0].cells[2]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

        for j in range(sample5.shape[-1]):
            col = table.columns[j].cells
            p=col[0].add_paragraph(sample5.columns[j])
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
            # table.cell(0,j).text = sample5.columns[j]
            
        for i in range(sample5.shape[0]):
            row = table.rows[i+1].cells
            for j in range(sample5.shape[-1]):
                p=row[j].add_paragraph(str(sample5.values[i,j]))
                p.alignment=WD_TABLE_ALIGNMENT.CENTER
                # table.cell(i+1,j).text = str(sample5.values[i,j])

                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), rgb_dict.get(sample5.values[i,0])))

                table.cell(i+1,j)._tc.get_or_add_tcPr().append(shading_elm_2)
    except:
        doc.add_paragraph('')
    ############## for adding hyperlink below the table ###############
    p = doc.add_paragraph('To download complete Inconsistencies Prefixes Data ')

    hyperlink = add_hyperlink(p, 'http://apihub.piloggroup.com:6652/user_down/inconsistencies_Prefixes/{}/{}/Inconsistencies Prefixes {}'.format(BATCH_ID, insert_time, format(datetime.now().replace(microsecond=0))), 'click here', '0000FF', True)

    ###########################################################################################

    # 4.3.1 Report on UOM's in Prefixes
    try:
        line006 = doc.add_paragraph().add_run('4.3.1 Detailed Data Analysis on Data Consistency')
        line006.font.size = Pt(14)
        line006.bold = True

        def non_standardized_uom_in_description():
            def uom_checker(description):
                container = []
                reference_list = {
                    "INCH": ["INCHES", "INCH", "IN" '"'],
                    "Current": ["AMPERES", "AMPERE", "AMPS", "AMP","A"],
                    "Voltage": ["VOLTS", "VOLT","V"],
                    "Meter": ["METERS", "METER", "MTR","M"]
                }
                for key, value in reference_list.items():
                    for item_inner in value:
                        pat = re.compile(r"\b[0-9]+[ ]?%s\b" % item_inner)
                        # pat = re.compile(r"\b[0-9]+[ ]?[^,]*%s\b" % item_inner)
                        try:
                            match = re.search(pat, description)
        #                     print(match)
                        except TypeError:
                            continue
                        if match:
                            container.append(item_inner + "+" + key)
                if container:
                    
                    return container
                else:
                    return None
            sample_df_4 = sample_df.drop_duplicates(subset=["LONG_DESCRIPTION"])
            sample_df_4["Captured_uom_list"] = sample_df_4["LONG_DESCRIPTION"].apply(lambda x: uom_checker(x))
            
            sample_df_4 = sample_df_4.dropna(subset=["Captured_uom_list"])
            
            sample_df_4_explode = sample_df_4.assign(Variant_Format=sample_df_4.Captured_uom_list).explode('Variant_Format')
            
            sample_df_4_explode["Standard UOM"] = sample_df_4_explode["Variant_Format"].apply(lambda x: x.split("+")[1])
            
            sample_df_4_explode["Variant UOM"] = sample_df_4_explode["Variant_Format"].apply(lambda x: x.split("+")[0])

            ############## for Number of Materials Linked ##########
            samp = sample_df_4_explode["Variant UOM"]
            sam_df = pd.DataFrame(samp)

            sam_df['MATERIAL']=sample_df['MATERIAL']
        
            def mat_Count(i):
                return sam_df["Variant UOM"].value_counts()[i]
            sam_df['Number of Materials Linked']=sam_df["Variant UOM"].apply(lambda x:mat_Count(x))
            sam_df.drop(columns=["Variant UOM"], inplace=True)
            ##################################
            sample_df_4_explode.drop_duplicates(subset=["Variant UOM"], inplace=True)

            ############
            frames = [sam_df,sample_df_4_explode]
            result = pd.concat(frames,axis=1)
            #####################
            list_to_concat_df = []
            for item in result["Standard UOM"].unique().tolist():
                if len(result[result["Standard UOM"] == item]["Variant UOM"].unique().tolist()) > 1:
                    list_to_concat_df.append(result[result["Standard UOM"] == item])   
            if list_to_concat_df:
                sample_df_4_output = pd.concat(list_to_concat_df)
                sample_df_4_output.drop(columns=["LONG_DESCRIPTION","MATERIAL","Variant_Format","Captured_uom_list"], inplace=True)
                sample_df_4_output.sort_values(by=["Variant UOM"], inplace=True)
                return sample_df_4_output

            else:
                return None
        sample4 = non_standardized_uom_in_description()
        sample44 = sample4.copy()


        # def UOM_Count(i):
        #     return sample4["Standard UOM"].value_counts()[i]
        # sample4['Number of Materials Linked']=sample4["Standard UOM"].apply(lambda x:UOM_Count(x))
        # def UOM_Counts(i):
        #     return sample44["Standard UOM"].value_counts()[i]
        # sample44['Number of Materials Linked']=sample44["Standard UOM"].apply(lambda x:UOM_Counts(x))




        ############# TO UPDATE DATA INTO DATABASE ##############

        sample44["BATCH_ID"] = BATCH_ID

        sample44["TABLE_TYPE"] = "UOM_Prefixes"
        sample44['INSERT_BY'] = '{}'.format(insert_time)

        sample44 = sample44.assign(VAL_SEQUENCE=[1 + i for i in range(len(sample44))])[['VAL_SEQUENCE'] + sample44.columns.tolist()]

        sample44 = sample44[['Standard UOM','Variant UOM','Number of Materials Linked','BATCH_ID','TABLE_TYPE','VAL_SEQUENCE', 'INSERT_BY']]

        sql = """insert into CLOUD_DHA(STANDARD_UOM,VARIANT_UOM,NUMBER_OF_MATERIALS_LINKED,BATCH_ID,TABLE_TYPE,VAL_SEQUENCE, INSERT_DATE ) values(:1, :2, :3, :4, :5, :6, :7)"""

        cursor = conn.cursor()
        for c, r in sample44.iterrows():
            cursor.execute(sql, tuple(r))
            conn.commit()
        # conn.close()

        #########################################################

        sample4 = sample4[['Standard UOM','Variant UOM','Number of Materials Linked']]
        rgb_list = ["E6E6FA", "F0FFF0", "FFF5EE", "EEE8AA", "7FFFD4", "696969", "E0FFFF","ADD8E6","87CEEB","5F9EA0","B0E0E6","AFEEEE","40E0D0","66CDAA","F08080","FFA07A","EE82EE","FFE4C4","FFFFE0","DCDCDC","F8F8FF"]
        unique_desc = sample4["Standard UOM"].unique().tolist()
        rgb_dict = dict(zip(unique_desc, rgb_list))
        table = doc.add_table(sample4.shape[0]+1,sample4.shape[1],style = 'Table Grid')
        row = table.rows[0].cells[2]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

        for j in range(sample4.shape[-1]):
            col = table.columns[j].cells
            p=col[0].add_paragraph(sample4.columns[j])
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
            # table.cell(0,j).text = sample5.columns[j]
            
        for i in range(sample4.shape[0]):
            row = table.rows[i+1].cells
            for j in range(sample4.shape[-1]):
                p=row[j].add_paragraph(str(sample4.values[i,j]))
                p.alignment=WD_TABLE_ALIGNMENT.CENTER
                # table.cell(i+1,j).text = str(sample5.values[i,j])

                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), rgb_dict.get(sample4.values[i,0])))

                table.cell(i+1,j)._tc.get_or_add_tcPr().append(shading_elm_2)
    except:
        doc.add_paragraph('')
    ############## for adding hyperlink below the table ###############
    
    p = doc.add_paragraph('To download complete UOM Prefixes Data ')
    hyperlink = add_hyperlink(p, 'http://apihub.piloggroup.com:6652/user_down/UOM_Prefixes/{}/{}/UOM Prefixes {}'.format(BATCH_ID, insert_time, format(datetime.now().replace(microsecond=0))), 'click here', '0000FF', True)
    
    ###########################################################################################

    # 4.3.2 Inconsistencies in the Description:
    try:
        line25 = doc.add_paragraph().add_run('4.3.2 Inconsistencies in the Description:')
        line25.bold = True
        line25.font.size = Pt(14)
        # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
        # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
        sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
        # sample_df = sample_df[:1500]
        # sample_df = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
        sample_df = sample_df.drop(columns=["UOM"])
        sample_df = sample_df.dropna(subset=["LONG_DESCRIPTION"])
        sample_df['LONG_DESCRIPTION'] = sample_df['LONG_DESCRIPTION'].apply(str)

        part_prefix_df = pd.read_excel(r'C:\Users\Administrator\Desktop\V10_ORIG.xls')
        part_prefix_df = part_prefix_df.fillna('')
        part_prefix_df = part_prefix_df.drop([37,40,41,47,48,58])
        def inconsistencies_in_the_description():
            def consistency_checker(description):
                container = []
                for i, row in part_prefix_df.iterrows():
                    if row.V10_PART_PREFIX.lower() != "manufacturer":
                        cleaned_part_code = re.sub(":", "", row["ORIGINAL_PART_PREFIX"])
                        cleaned_part_code = re.sub("\.", "\\.", cleaned_part_code)
                        pat = re.compile(r"(\b%s.?[:].....[^, ]*)" % cleaned_part_code)
                        # pat = re.compile(r"(\b%s.?[:].....[^, ]*)" % cleaned_part_code)
                        try:
                            match = re.search(pat, description)
                        except TypeError:
                            continue
                        if match:
                            captured_match = row.V10_PART_PREFIX + ":" + match.group()
                            if captured_match not in container:
                                container.append(captured_match)
                if container:
                    return container
                else:
                    return None
                
            sample_df_2 =sample_df["LONG_DESCRIPTION"].apply(lambda x: re.sub(r" :",":",x))
            sample_2_df = pd.DataFrame(sample_df_2)
            sample_2_df = sample_2_df.drop_duplicates(subset=["LONG_DESCRIPTION"])
            sample_2_df["Captured_part_list"] = sample_2_df["LONG_DESCRIPTION"].apply(lambda x: consistency_checker(x))
            sample_2_df = sample_2_df.dropna(subset=["Captured_part_list"])
            sample_2_df_explode = sample_2_df.assign(Captured_parts=sample_2_df.Captured_part_list).explode('Captured_parts')
            sample_2_df_explode["Captured_V010"] = sample_2_df_explode["Captured_parts"].apply(lambda x: x.split(":")[0])
            sample_2_df_explode["Captured_Pre"] = sample_2_df_explode["Captured_parts"].apply(lambda x: x.split(":")[1])
            #############
            samp = sample_2_df_explode["Captured_Pre"]
            sam_df = pd.DataFrame(samp)
            def mat_Count(i):
                return sam_df["Captured_Pre"].value_counts()[i]
            sam_df['Number of Materials Linked']=sam_df["Captured_Pre"].apply(lambda x:mat_Count(x))
            sam_df.drop(columns=["Captured_Pre"], inplace=True)
            ################

            sample_2_df_explode["Captured_Post"] = sample_2_df_explode["Captured_parts"].apply(lambda x: x.split(":")[2])
            sample_2_df_explode.drop(columns=["Captured_part_list", "Captured_parts", "Captured_Post"], inplace=True)
            sample_2_df_explode.drop_duplicates(subset=["Captured_Pre"],inplace=True)

            #####@@@@@@
            sam_df.reset_index(drop=True, inplace=True)
            sample_2_df_explode.reset_index(drop=True, inplace=True)
            result = pd.concat([sam_df, sample_2_df_explode], axis=1)
            ####@@@@@@@@
            
            list_to_concat_df = []
            for item in result["Captured_V010"].unique().tolist():
                if len(result[result["Captured_V010"] == item]["Captured_Pre"].unique().tolist()) > 1:
                    list_to_concat_df.append(result[result["Captured_V010"] == item])


            if list_to_concat_df:
                sample_2_df_output = pd.concat(list_to_concat_df)
                sample_2_df_output["Remarks"] = sample_2_df_output.apply(
                    lambda x: "The format of %s is not standardized." % (x["Captured_V010"]), axis=1)

                sample_2_df_output.drop(columns=["Captured_V010"], inplace=True)
                sample_2_df_output.sort_values(by=["Remarks"], inplace=True)

                sample_2_df_output['MATERIAL']=sample_df['MATERIAL']
                cols = sample_2_df_output.columns.tolist()
                cols = cols[-1:] + cols[:-1]
                sample_2_df_output = sample_2_df_output[cols]
                return sample_2_df_output
            else:
                return None

        sample2 = inconsistencies_in_the_description()
        sample2.to_excel(r'ckeck1.xlsx')

        sample22 = sample2.copy()

        sample22.drop(columns=["Number of Materials Linked"], inplace=True)



        sample12 = sample2.copy()
        sample12.to_excel(r'check2.xlsx')
        # print(sample12)

        # def mat_Count(i):
        #     return sample12["Standard Format"].value_counts()[i]

        # sample12['Number_of_Materials_Linked']=sample12["Standard Format"].apply(lambda x:mat_Count(x))

        # print(sample12,'=+++++++++++++++++++++++++++++=')

        # sample12.sort_values(by=["Standard Format"], inplace=True)
        # sample12.drop(columns=["Captured_Pre"], inplace=True)
        # sample12.to_excel(r'ckeck2.xlsx')





        # ########### UPDATE DATA INTO DATABASE #############

        sample12["BATCH_ID"] = BATCH_ID

        sample12["TABLE_TYPE"] = "inconsistencies"
        
        sample12['INSERT_BY'] = '{}'.format(insert_time)

        sample12 = sample12.assign(VAL_SEQUENCE=[1 + i for i in range(len(sample12))])[['VAL_SEQUENCE'] + sample12.columns.tolist()]

        sample12 = sample12[['MATERIAL','LONG_DESCRIPTION','Captured_Pre','Remarks','BATCH_ID','Number of Materials Linked','TABLE_TYPE','VAL_SEQUENCE', 'INSERT_BY']]
        sample12.to_excel(r'check3.xlsx')

        sql = """insert into CLOUD_DHA(MATERIAL,LONGDESCRIPTION,CAPTURED_PRE,REMARKS,BATCH_ID,NUMBER_OF_MATERIALS_LINKED,TABLE_TYPE,VAL_SEQUENCE, INSERT_DATE ) values(:1, :2, :3, :4, :5, :6, :7, :8, :9)"""

        cursor = conn.cursor()
        for c, r in sample12.iterrows():
            print('--------------------->',r)
            cursor.execute(sql, tuple(r))
            conn.commit()




        # # conn.close()

        # #####################################################

        list_to_color = sample22["Captured_Pre"].tolist()
        sample22 = sample22.drop(columns=["Captured_Pre"])


        table = doc.add_table(sample22.shape[0]+1, sample22.shape[1],style = 'Table Grid')
        row=table.rows[0].cells[2]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
        table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
        table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

        # lst = [914400,3657600,914400]
        for j in range(sample22.shape[-1]):

            # table.cell(0,j).width = lst[j]
            # table.cell(0,j).text = sample2.columns[j]
            col = table.columns[j].cells
            p=col[0].add_paragraph(sample22.columns[j])
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
            
            
        for i in range(sample22.shape[0]):
            row = table.rows[i+1].cells
            for j in range(sample22.shape[-1]):

                # table.cell(i+1,j).width = lst[j]

                p=row[j].add_paragraph(str(sample22.values[i,j]))
                p.alignment=WD_TABLE_ALIGNMENT.CENTER
                # table.cell(i+1,j).text = str(sample2.values[i,j])
                if j == 1:
                    list_to_color_specific = list_to_color[i].split(",")
                    list_to_color_specific = [re.sub(r"\s+$", "", re.sub(r"^\s+", "", x)) for x in list_to_color_specific]
                    for paragraph in table.cell(i+1,j).paragraphs:
                        desc = paragraph.text
                        paragraph.clear()
                        # desc = desc.replace('OEMPART NO','11111')
                        # for word in re.split(r"\b", desc):
                            # if word == '11111':
                            #         word = 'OEMPART NO'
                            # print(word)
                        # for l in list_to_color_specific:
                        #     if re.search(l,desc):

                        #     # if word in list_to_color_specific:
                        #         run = paragraph.add_run()
                        #         # run.text = word
                        #         run.text = l
                        #         run.font.color.rgb = RGBColor(0xFF, 0x11, 0x11)
                        #     else:
                        #         run = paragraph.add_run()
                        #         # run.text = word
                        #         run.text = l
                        for l in list_to_color_specific:
                            if re.search(l,desc):
                                desc = re.sub(l," {} ".format(l), desc)
                                run = paragraph.add_run()
                                run.text = desc.split(l)[0]
                                run = paragraph.add_run()
                                run.text = l
                                run.font.color.rgb = RGBColor(0xFF, 0x11, 0x11)
                                run = paragraph.add_run()
                                run.text = desc.split(l)[1]
                                break

    except:
        doc.add_paragraph('')
    ############## for adding hyperlink below the table ###############

    p = doc.add_paragraph("To download complete Inconsistencies Data ")

    hyperlink = add_hyperlink(p, 'http://apihub.piloggroup.com:6652/user_down/inconsistencies/{}/{}/Inconsistencies {}'.format(BATCH_ID, insert_time, format(datetime.now().replace(microsecond=0))), 'click here', '0000FF', True)


    ###########################################################################################


    doc.add_page_break()
    # 4.3.3 Non-standardized UOMs in Description
    try:
        line29 = doc.add_paragraph().add_run('4.3.3 Non-standardized UOMs in Description')
        line29.bold = True
        line29.font.size = Pt(14)
        def non_standardized_uom_in_description():
            def uom_checker(description):
                container = []
                reference_list = {
                    "length": ["INCHES", "INCH", "IN"],
                    "current": ["AMPERES", "AMPERE", "AMPS", "AMP"],
                    "voltage": ["VOLTS", "VOLT"],
                    "met_length": ["METERS", "METER", "MTR"]
                }
                for key, value in reference_list.items():
                    for item_inner in value:
                        pat = re.compile(r"\b[0-9]+[ ]?%s\b" % item_inner)
                        # pat = re.compile(r"\b[0-9]+[ ]?[^,]*%s\b" % item_inner)
                        try:
                            match = re.search(pat, description)
                        except TypeError:
                            continue
                        if match:
                            container.append(item_inner + "+" + key)
                if container:
                    
                    return container
                else:
                    return None

            sample_df_4 = sample_df.drop_duplicates(subset=["LONG_DESCRIPTION"])
            sample_df_4["Captured_uom_list"] = sample_df_4["LONG_DESCRIPTION"].apply(lambda x: uom_checker(x))
            sample_df_4 = sample_df_4.dropna(subset=["Captured_uom_list"])
            sample_df_4_explode = sample_df_4.assign(Captured_uom=sample_df_4.Captured_uom_list).explode('Captured_uom')
            sample_df_4_explode["Uom_standard"] = sample_df_4_explode["Captured_uom"].apply(lambda x: x.split("+")[1])
            sample_df_4_explode["Captured_uom"] = sample_df_4_explode["Captured_uom"].apply(lambda x: x.split("+")[0])
            sample_df_4_explode.drop_duplicates(subset=["Captured_uom"], inplace=True)

            list_to_concat_df = []
            for item in sample_df_4_explode["Uom_standard"].unique().tolist():
                if len(sample_df_4_explode[sample_df_4_explode["Uom_standard"] == item]["Captured_uom"].unique().tolist()) > 1:
                    list_to_concat_df.append(sample_df_4_explode[sample_df_4_explode["Uom_standard"] == item])
            if list_to_concat_df:
                sample_df_4_output = pd.concat(list_to_concat_df)
                sample_df_4_output = sample_df_4_output.groupby(["MATERIAL", "LONG_DESCRIPTION"], as_index=False).agg(Captured_uom=("Captured_uom", ",".join))
                sample_df_4_output["Remarks"] = "UOM's Captured in Different Format"
                sample_df_4_output.sort_values(by=["Captured_uom"], inplace=True)
                # sample_df_4_output = sample_df_4_output.drop(columns=["Captured_uom"])
                return sample_df_4_output

            else:
                return None


        sample4 = non_standardized_uom_in_description()
        print(sample4)

        sample14 = sample4.copy()
        # print('=========',sample14)

        # ########  DATA UPDATE INTO DATABASE  ############
        sample14["BATCH_ID"] = BATCH_ID

        sample14["TABLE_TYPE"] = "non_standardized_uom"
        
        sample14['INSERT_BY'] = '{}'.format(insert_time)

        sample14 = sample14.assign(VAL_SEQUENCE=[1 + i for i in range(len(sample14))])[['VAL_SEQUENCE'] + sample14.columns.tolist()]

        sample14 = sample14[['MATERIAL','LONG_DESCRIPTION','Captured_uom','Remarks','BATCH_ID','TABLE_TYPE','VAL_SEQUENCE', 'INSERT_BY']]


        sql = """insert into CLOUD_DHA(MATERIAL,LONGDESCRIPTION,CAPTURED_PRE,REMARKS,BATCH_ID,TABLE_TYPE,VAL_SEQUENCE, INSERT_DATE ) values(:1, :2, :3, :4, :5, :6, :7, :8 )"""

        cursor = conn.cursor()
        for c, r in sample14.iterrows():
            cursor.execute(sql, tuple(r))
            conn.commit()
        # conn.close()

        # ##################################################


        list_to_color = sample4["Captured_uom"].tolist()
        sample4 = sample4.drop(columns=["Captured_uom"])
        # Sample4=pd.DataFrame(sample_df_4_output)

        table = doc.add_table(sample4.shape[0]+1, sample4.shape[1],style = 'Table Grid')
        row=table.rows[0].cells[2]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#85bae5"/>'.format(nsdecls('w')))
        table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

        lst = [914400,3657600,914400]

        for j in range(sample4.shape[-1]):
            table.cell(0,j).width = lst[j]
            # table.cell(0,j).text = sample4.columns[j]
            col = table.columns[j].cells
            p=col[0].add_paragraph(sample4.columns[j])
            p.alignment=WD_TABLE_ALIGNMENT.CENTER

        for i in range(sample4.shape[0]):
            row = table.rows[i+1].cells
            for j in range(sample4.shape[-1]):
                table.cell(i+1,j).width = lst[j]
                p=row[j].add_paragraph(str(sample4.values[i,j]))
                p.alignment=WD_TABLE_ALIGNMENT.CENTER
                # table.cell(i+1,j).text = str(sample4.values[i,j])
                if j == 1:
                    list_to_color_specific = list_to_color[i].split(",")
                    list_to_color_specific = [re.sub(r"\s+$", "", re.sub(r"^\s+", "", x)) for x in list_to_color_specific]
                    for paragraph in table.cell(i+1,j).paragraphs:
                        desc = paragraph.text
                        paragraph.clear()
                        for word in re.split(r"\b", desc):
                            if re.sub(r"^\d+", "", word) in list_to_color_specific:
                                run = paragraph.add_run()
                                run.text = re.findall(r"^\d+", word)
                                run = paragraph.add_run()
                                run.text = re.findall(re.sub(r"^\d+", "", word), word)
                                run.font.color.rgb = RGBColor(0xFF, 0x11, 0x11)
                            else:
                                if word=='",':
                                    run = paragraph.add_run()
                                    run.text = '"'
                                    run.font.color.rgb = RGBColor(0xFF, 0x11, 0x11)
                                    run = paragraph.add_run()
                                    run.text = ','
                                else:
                                    run = paragraph.add_run()
                                    run.text = word
    except:
        doc.add_paragraph('')
    ############## for adding hyperlink below the table ###############
    p = doc.add_paragraph('To download complete Non-Standardized UOMs Data ')

    hyperlink = add_hyperlink(p, 'http://apihub.piloggroup.com:6652/user_down/non_standardized_uom/{}/{}/Non-Standardized UOMs {}'.format(BATCH_ID, insert_time, format(datetime.now().replace(microsecond=0))), 'click here', '0000FF', True)

    ###########################################################################################

    # 5.4 Analysis on Top 10 UOM'S
    doc.add_page_break()
    line31 = doc.add_paragraph().add_run('4.4 Analysis on Top 10 UOMs')
    line31.bold = True
    line31.font.size = Pt(16)
    line32 = doc.add_paragraph().add_run('UOMs are listed below:')
    # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\keshav_sir_files\DHA3.xlsx')
    # sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA1.xlsx')[:1500]
    # sample_df = pd.read_sql("select * from O_RECORD_HEALTH_ASSESSMENT where BATCH_ID = '{}' ".format(BATCH_ID), conn)
    sample_df = pd.read_excel(r'C:\Users\Administrator\Desktop\all files\madam_files\DHA_DATA2.xlsx')[:1000]
    sample_df = sample_df.fillna('')
    UOM = sample_df['UOM'].value_counts()
    Top_10 = UOM.head(10)
    # Top_10
    UOM1 = pd.DataFrame(Top_10)
    UOM1.reset_index(level=[0], drop=False, inplace=True)
    UOM1.rename(columns = {"index": "UOM","UOM":'Top10 Count'}, inplace = True)
    fig = px.pie(UOM1, values='Top10 Count',names='UOM',color_discrete_sequence=["rgb(168, 162, 64)", "rgb(64, 135, 94)", "rgb(240, 171, 0)", "rgb(0, 113, 197)", "rgb(255,0,255)", "rgb(148,0,211)", "rgb(210,105,30)", "rgb(0,255,255)","rgb(255,127,80)", "rgb(128,0,0)"])
    fig.update_traces(textposition='inside',textinfo='value')
    fig.update_layout(  title={'text':'Top 10 UOMS','y':0.95, 'x':0.5,'xanchor':'center','yanchor':'top'},
                        titlefont=dict( family='Times New Roman',
                                        size=30,
                                        color='Red'
                                        ),
                        xaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'white',
                            zerolinecolor= 'Black'
                            #domain=[0.15, 1]
                        ),
                yaxis=dict(
                            showgrid=True,
                            showline=True,
                            showticklabels=True,
                            zeroline=True,
                            linecolor='Black',
                            linewidth=2,
                            gridcolor= 'White',
                            zerolinecolor= 'Black'
                        ),
                            showlegend=True,
                            paper_bgcolor='rgb(248, 248, 255)',
                            plot_bgcolor='rgb(248, 246, 255)',
                            margin=dict(l=100, r=20, t=70, b=70),
                        uniformtext_minsize=20, uniformtext_mode='hide')
    config={'displaylogo':False}

    table = doc.add_table(UOM1.shape[0]+1, UOM1.shape[1],style = 'Table Grid')
    row = table.rows[0].cells[1]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#76b5c5"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)

    for j in range(UOM1.shape[-1]):
        col = table.columns[j].cells
        p=col[0].add_paragraph(UOM1.columns[j])
        p.alignment=WD_TABLE_ALIGNMENT.CENTER

        # table.cell(0,j).text = UOM1.columns[j]
        # table.cell(0.j).text.alignment =WD_TABLE_ALIGNMENT.CENTER 

    for i in range(UOM1.shape[0]):
        row = table.rows[i+1].cells
        for j in range(UOM1.shape[-1]):
            p=row[j].add_paragraph(str(UOM1.values[i,j]))
            p.alignment=WD_TABLE_ALIGNMENT.CENTER
            # table.cell(i+1,j).text = str(UOM1.values[i,j])


    fig.write_image(r'Top10_UOM.png')
    doc.add_paragraph('\n')
    doc.add_picture(r'Top10_UOM.png',width=Inches(6), height=Inches(3))
    # doc.add_page_break()


    #5. Key Recommendations
    line17 = doc.add_paragraph().add_run('5. Key Recommendations')
    line17.bold = True
    line17.font.size = Pt(18)
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Establish a Master Data vision that recognizes Master Data as an asset that reduces cost, risk and exposure through efficient use of unique and reliable master data within the organization, related to external partners and customers')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('The organization should convene a working group (e.g., data stewards) representing all relevant stakeholders to determine targets, set thresholds, and define the quality dimensions that are most important')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Periodic assessments should be conducted to determine if acceptable thresholds and targets are being met, and metrics should be updated accordingly')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Implement data standards to drive standard data in accordance with data requirements')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Implement data workflows to manage data approval criteria where the engineers review the data quality')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Implement advanced Master Data governance tools to control the implementation of data standards, work flows, enrichment of data, data lifecycle and uniqueness of data')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Set up analytics for contracting and spend analysis')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Establish Business Rules to align the organization on Cataloguing Standards such as language, measurements, units of measure, presentation of data, etc. using ISO 8000 standards')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Create standardized Short and Long Descriptions for all items to ensure purchasing accuracy')
    intt = doc.add_paragraph(style='List Bullet 2').add_run('Harmonize data with a final potential duplicate resolution process across and languages to ensure excess stock is used and not reordered or duplicated between the sites')




    for paragraph in doc.paragraphs:
        if 'SAUDI CEMENT' in paragraph.text:
            inline=paragraph.runs
            for i in range(len(inline)):
                if 'SAUDI CEMENT' in inline[i].text:
                    text=inline[i].text.replace('SAUDI CEMENT','Organization')
                    inline[i].text=text


    # doc_unq_name = format(datetime.now().replace(microsecond=0))

    doc.save(r'DHA_02_04_22.docx')
    

    # inputFile = os.path.abspath(r"DHA_DOC3.docx")
    # outputFile = os.path.abspath(r"DHA_DOC3.pdf")
    # word = win32com.client.Dispatch('Word.Application')
    # doc = word.Documents.Open(inputFile)
    # doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
    # doc.Close()
    # word.Quit()


    # # convert("DHA_DOC3.docx", "DHA_DOC3.pdf")

    # # convert_to_pdf_win(r'C:/Users/Administrator/Desktop/FEB_2022/DHA_DOC3.docx', 'C:/Users/Administrator/Desktop/FEB_2022/')
    # # convert_to_pdf_win(r'C:/Users/Administrator/Desktop/FEB_2022/DHA_DOC3.docx', 'C:/Users/Administrator/Desktop/FEB_2022/')
    # # convert(r"DHA_DOC3.docx")

    # pdf_file = "C://Users//Administrator//Desktop//FEB_2022//DHA_DOC3.pdf"
    # watermark = "C://Users//Administrator//Desktop//NOV//Wmark.pdf"
    # merged = "C://Users//Administrator//Desktop//FEB_2022//DHA_OUTPUT_01_01_22.pdf"

    # with open(pdf_file, "rb") as input_file, open(watermark, "rb") as watermark_file:
    #     input_pdf = PdfFileReader(input_file)
    #     watermark_pdf = PdfFileReader(watermark_file)
    #     watermark_page = watermark_pdf.getPage(0)

    #     output = PdfFileWriter()

    #     for i in range(input_pdf.getNumPages()):
    #         pdf_page = input_pdf.getPage(i)
    #         pdf_page.mergePage(watermark_page)
    #         output.addPage(pdf_page)

    #     with open(merged, "wb") as merged_file:
    #         output.write(merged_file)


    END = datetime.now()
    print('End time',END)

    return FileResponse(r'DHA_02_04_22.docx')

END = datetime.now()
print('End time',END)

# download_dct = {'inconsistencies' : "MATERIAL,LONGDESCRIPTION,CAPTURED_PRE,REMARKS,NUMBER_OF_MATERIALS_LINKED", 
#                 'Reference_Details' : "MATERIAL,LONGDESCRIPTION,EXTRACTED_REFERENCE_NUMBERS,REFERENCE_NUMBER_PREFIX_TYPE",
#                 # 'inconsistencies_Prefixes' : "STANDARD_FORMAT, VARIANT_FORMAT, NUMBER_OF_MATERIALS_LINKED",
#                 'UOM_Prefixes' : "STANDARD_UOM, VARIANT_UOM, NUMBER_OF_MATERIALS_LINKED",
#                 'non_standardized_uom' : "MATERIAL, CAPTURED_PRE, LONGDESCRIPTION, REMARKS"}



# def ref_name(tabletype, batch_id, inst_time, time_stamp):
# # download_dct = {'inconsistencies' : "MATERIAL,LONGDESCRIPTION,CAPTURED_PRE,REMARKS"}
#     print(download_dct[tabletype])
#     conn = cx_Oracle.connect("DR1024193/Pipl#mdrm$93@172.16.1.61/DR101412")
#     s_qry = """select {} from CLOUD_DHA where BATCH_ID = '{}' and TABLE_TYPE = '{}' and INSERT_DATE = '{}' """.format(download_dct[tabletype], batch_id, tabletype, inst_time)
#     print(s_qry)
#     fetch = pd.read_sql(s_qry,conn)
#     reg = fetch.to_excel("User_Download.xlsx")
#     conn.close()
#     return "User_Download.xlsx"

# def ref_name():
    # conn = cx_Oracle.connect("DR1024193/Pipl#mdrm$93@172.16.1.61/DR101412")
    # fetch = pd.read_sql("""select MATERIAL,LONGDESCRIPTION,CAPTURED_PRE,REMARKS from CLOUD_DHA where BATCH_ID ='00003218' and TABLE_TYPE ='inconsistencies' """,conn)
    # reg = fetch.to_excel("Download.xlsx")
    # conn.close()
    # return "Download.xlsx"

# app = FastAPI()


# @app.get("/user_down/{tabletype}/{batch_id}/{inst_time}/{time_stamp}")
# def main_download(tabletype, batch_id, inst_time, time_stamp):
#     return FileResponse(ref_name(tabletype, batch_id , inst_time, time_stamp))


if __name__ == '__main__':
    uvicorn.run(app, host="127.0.0.1", port= 5555 , loop='asyncio')